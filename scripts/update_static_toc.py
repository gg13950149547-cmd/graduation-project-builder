from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from collections import Counter
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ALLOWED_PACKAGE_DRIFT = {"word/document.xml"}
REQUIRED_TAIL_TOC_LABELS = {"\u53c2\u8003\u6587\u732e", "\u81f4\u8c22"}


def package_part_hashes(path: Path) -> dict[str, str]:
    hashes: dict[str, str] = {}
    with zipfile.ZipFile(path, "r") as zf:
        for name in zf.namelist():
            hashes[name] = hashlib.sha256(zf.read(name)).hexdigest()
    return hashes


def assert_only_allowed_package_drift(before_path: Path, after_path: Path) -> None:
    before = package_part_hashes(before_path)
    after = package_part_hashes(after_path)
    issues: list[str] = []
    before_names = set(before)
    after_names = set(after)
    for name in sorted(before_names - after_names):
        if name not in ALLOWED_PACKAGE_DRIFT:
            issues.append(f"removed package part outside TOC surface: {name}")
    for name in sorted(after_names - before_names):
        if name not in ALLOWED_PACKAGE_DRIFT:
            issues.append(f"added package part outside TOC surface: {name}")
    for name in sorted(before_names & after_names):
        if before[name] != after[name] and name not in ALLOWED_PACKAGE_DRIFT:
            issues.append(f"changed package part outside TOC surface: {name}")
    if issues:
        raise RuntimeError(
            "static TOC update changed package parts outside word/document.xml; "
            "refusing unsafe python-docx roundtrip: " + "; ".join(issues)
        )


def body_direct_paragraphs(root):
    body = root.find(qn("w:body"))
    if body is None:
        raise RuntimeError("document.xml has no body")
    return [child for child in list(body) if child.tag == qn("w:p")]


def replace_body_paragraph(parent_body, old_paragraph, new_paragraph) -> None:
    children = list(parent_body)
    child_index = children.index(old_paragraph)
    parent_body.remove(old_paragraph)
    parent_body.insert(child_index, deepcopy(new_paragraph))


def keep_only_toc_document_xml_delta(before_path: Path, after_path: Path, toc_indexes: set[int]) -> None:
    temp_path = after_path.with_suffix(after_path.suffix + ".tmp")
    with zipfile.ZipFile(after_path, "r") as after_zf:
        updated_document_xml = after_zf.read("word/document.xml")
    with zipfile.ZipFile(before_path, "r") as before_zf:
        original_document_xml = before_zf.read("word/document.xml")

    original_root = etree_from_bytes(original_document_xml)
    updated_root = etree_from_bytes(updated_document_xml)
    original_body = original_root.find(qn("w:body"))
    if original_body is None:
        raise RuntimeError("input document.xml has no body")
    original_paragraphs = body_direct_paragraphs(original_root)
    updated_paragraphs = body_direct_paragraphs(updated_root)
    if len(original_paragraphs) != len(updated_paragraphs):
        raise RuntimeError(
            "static TOC update changed top-level paragraph count; refusing broad document.xml roundtrip"
        )
    for idx in sorted(toc_indexes):
        if idx < 0 or idx >= len(original_paragraphs):
            raise RuntimeError(f"static TOC paragraph index out of range after save: {idx}")
        replace_body_paragraph(original_body, original_paragraphs[idx], updated_paragraphs[idx])
        original_paragraphs[idx] = updated_paragraphs[idx]
    patched_document_xml = serialize_xml(original_root)
    with zipfile.ZipFile(before_path, "r") as before_zf, zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as out_zf:
        for info in before_zf.infolist():
            data = patched_document_xml if info.filename == "word/document.xml" else before_zf.read(info.filename)
            out_zf.writestr(info, data)
    temp_path.replace(after_path)


def etree_from_bytes(payload: bytes):
    from xml.etree import ElementTree as ET

    return ET.fromstring(payload)


def serialize_xml(root) -> bytes:
    from xml.etree import ElementTree as ET

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def normalize(text: str) -> str:
    return re.sub(r"[\s\u3000\u25a1]+", "", text or "")


def normalize_heading_dots(text: str) -> str:
    return re.sub(r"(?<=\d)[\s\u25a1]*[\.\uff0e][\s\u25a1]*(?=\d)", ".", text or "")


def heading_key(text: str) -> str:
    head = re.split(r"[\(\uff08]", str(text or "").strip(), maxsplit=1)[0]
    return normalize(head)


def is_toc_heading_text(text: str) -> bool:
    normalized = normalize(text).lower()
    key = heading_key(text).lower()
    target = normalize("\u76ee\u5f55").lower()
    return (
        key in {target, "contents", "tableofcontents"}
        or normalized in {target, "contents", "tableofcontents"}
        or key.endswith(target)
        or normalized.endswith(target)
    )


def contains_chapter_heading_marker(text: str) -> bool:
    stripped = re.sub(r"^[\s\u25a1]+", "", text or "")
    match = re.match(r"^(\u7b2c[0-9\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u7ae0)(.*)$", stripped)
    if not match:
        return False
    tail = match.group(2)
    if not tail:
        return True
    separator_match = re.match(r"^[\s\u25a1:：、.\-—]+", tail)
    if separator_match:
        tail = tail[separator_match.end() :]
    compact_tail = normalize(tail)
    sentence_punctuation = "\uff0c\uff1b\uff1a\uff01\uff1f\u3002\uff08\uff09,;:!?()"
    if any(ch in compact_tail for ch in sentence_punctuation):
        return False
    return len(compact_tail) <= (24 if separator_match else 12)


def heading_level(text: str) -> int | None:
    stripped = re.sub(r"^[\s\u25a1]+", "", text.strip())
    sep = r"[\s\u25a1]+"
    normalized_dots = normalize_heading_dots(stripped)
    if re.match(rf"^\d{{1,2}}{sep}\S", stripped) or contains_chapter_heading_marker(stripped):
        return 1
    if re.match(rf"^\d{{1,2}}\.\d+{sep}\S", normalized_dots):
        return 2
    if re.match(rf"^\d{{1,2}}\.\d+\.\d+{sep}\S", normalized_dots):
        return 3
    if re.match(rf"^\d{{1,2}}\.\d+\.\d+\.\d+{sep}\S", normalized_dots):
        return 4
    return None


def set_text_node(text_node, value: str) -> None:
    text_node.text = value
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"
    if value.startswith(" ") or value.endswith(" "):
        text_node.set(xml_space, "preserve")
    elif xml_space in text_node.attrib:
        del text_node.attrib[xml_space]


def replace_trailing_page_number(text: str, page: str) -> str | None:
    """Replace only a recognizable TOC page-number tail, never the whole entry."""
    value = text or ""
    if not value.strip():
        return None
    # A no-tab static TOC entry is safe to patch only when the page number is
    # visibly separated from the title by dotted leaders or a wide whitespace gap.
    match = re.search(r"(?P<prefix>.*(?:\u2026|\.{2,}|\s{2,})\s*)(?P<page>\d+|[ivxlcdmIVXLCDM]+)\s*$", value)
    if not match:
        return None
    prefix = match.group("prefix")
    if not prefix.strip():
        return None
    return prefix + page


def set_toc_page_preserving_runs(paragraph, page: str) -> None:
    seen_tab = False
    page_written = False
    last_run = None
    text_nodes = []
    for run_obj in paragraph.runs:
        run_element = run_obj._element
        last_run = run_element
        for child in list(run_element):
            if child.tag == qn("w:tab"):
                seen_tab = True
                continue
            if child.tag != qn("w:t"):
                continue
            text_nodes.append(child)
            if not seen_tab:
                continue
            if not page_written:
                set_text_node(child, page)
                page_written = True
            else:
                set_text_node(child, "")
    if not seen_tab and text_nodes:
        for text_node in reversed(text_nodes):
            patched = replace_trailing_page_number(text_node.text or "", page)
            if patched is not None:
                set_text_node(text_node, patched)
                page_written = True
                break
        if not page_written:
            raise RuntimeError(
                "locked TOC entry lacks tab/leader/page-number run segmentation; "
                "refusing to replace the whole entry with a page number"
            )
    if not seen_tab:
        if not page_written:
            run_obj = paragraph.add_run()
            text_node = OxmlElement("w:t")
            set_text_node(text_node, page)
            run_obj._element.append(text_node)
    elif not page_written:
        run_obj = paragraph.add_run()
        if last_run is not None:
            rpr = last_run.find(qn("w:rPr"))
            if rpr is not None:
                run_obj._element.insert(0, deepcopy(rpr))
        text_node = OxmlElement("w:t")
        set_text_node(text_node, page)
        run_obj._element.append(text_node)


FONT_ATTRS = (
    "eastAsia",
    "ascii",
    "hAnsi",
    "cs",
    "eastAsiaTheme",
    "asciiTheme",
    "hAnsiTheme",
    "csTheme",
)
FONT_ALIAS_MAP = {
    "simsun": "\u5b8b\u4f53",
    "simhei": "\u9ed1\u4f53",
    "kaiti": "\u6977\u4f53",
    "fangsong": "\u4eff\u5b8b",
    "\u5b8b\u4f53": "\u5b8b\u4f53",
    "\u9ed1\u4f53": "\u9ed1\u4f53",
    "\u6977\u4f53": "\u6977\u4f53",
    "\u4eff\u5b8b": "\u4eff\u5b8b",
}


def preferred_font(value: str) -> str:
    parts = [part.strip() for part in str(value or "").split(";") if part.strip()]
    if not parts:
        return value
    for part in parts:
        mapped = FONT_ALIAS_MAP.get(part.lower()) or FONT_ALIAS_MAP.get(part)
        if mapped:
            return mapped
    for part in parts:
        if any(ord(ch) > 127 for ch in part):
            return part
    return parts[0]


def iter_run_elements(paragraph):
    return paragraph._element.findall(".//" + qn("w:r"))


def paragraph_has_run_tab(paragraph) -> bool:
    return "\t" in (paragraph.text or "") or any(
        run_element.find(qn("w:tab")) is not None for run_element in iter_run_elements(paragraph)
    )


def direct_toc_tab_stops(paragraph):
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        return []
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        return []
    return list(tabs.findall(qn("w:tab")))


def normalize_toc_right_tab_positions(doc: Document, allowed_toc_indexes: set[int]) -> int:
    positions: Counter[str] = Counter()
    tab_elements = []
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if para_idx not in allowed_toc_indexes:
            continue
        if not paragraph_has_run_tab(paragraph) and not is_toc_leader_entry(paragraph.text):
            continue
        for tab in direct_toc_tab_stops(paragraph):
            if (tab.get(qn("w:val")) or "left") != "right":
                continue
            if (tab.get(qn("w:leader")) or "none") != "dot":
                continue
            pos = tab.get(qn("w:pos")) or ""
            if not pos:
                continue
            positions[pos] += 1
            tab_elements.append(tab)
    if not positions:
        return 0
    preferred_pos = positions.most_common(1)[0][0]
    changed = 0
    for tab in tab_elements:
        if (tab.get(qn("w:pos")) or "") != preferred_pos:
            tab.set(qn("w:pos"), preferred_pos)
            changed += 1
    return changed


def normalize_toc_entry_paragraph_spacing(doc: Document, allowed_toc_indexes: set[int]) -> int:
    changed = 0
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if para_idx not in allowed_toc_indexes:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if style_name not in {"toc 1", "toc 2", "toc 3"} and not paragraph_has_run_tab(paragraph):
            continue
        ppr = paragraph._element.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            paragraph._element.insert(0, ppr)
        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            ppr.append(spacing)
        expected = {
            qn("w:before"): "0",
            qn("w:after"): "0",
            qn("w:lineRule"): "auto",
            qn("w:line"): "360",
        }
        for attr_name, value in expected.items():
            if spacing.get(attr_name) != value:
                spacing.set(attr_name, value)
                changed += 1
    return changed


def toc_paragraph_indexes(doc: Document) -> set[int]:
    start: int | None = None
    for idx, para in enumerate(doc.paragraphs):
        style_name = (para.style.name if para.style else "").lower()
        if is_toc_heading_text(para.text) or style_name == "toc heading":
            start = idx + 1
            break
    if start is None:
        raise RuntimeError("locked TOC range not found; refusing whole-document TOC rewrite")
    end = len(doc.paragraphs)
    for idx in range(start, len(doc.paragraphs)):
        para = doc.paragraphs[idx]
        style_name = (para.style.name if para.style else "").lower()
        if (
            heading_level(para.text) is not None
            and not style_name.startswith("toc ")
            and not paragraph_has_run_tab(para)
            and not is_toc_leader_entry(para.text)
        ):
            end = idx
            break
    indexes = {
        idx
        for idx in range(start, end)
        if doc.paragraphs[idx].text.strip()
        or (doc.paragraphs[idx].style and (doc.paragraphs[idx].style.name or "").lower().startswith("toc "))
    }
    if not indexes:
        raise RuntimeError("locked TOC range is empty; refusing whole-document TOC rewrite")
    return indexes


def require_mapped_tail_entries_present(doc: Document, allowed_toc_indexes: set[int], page_map: dict[str, str]) -> None:
    present = {
        normalize(toc_visible_label(paragraph.text))
        for para_idx, paragraph in enumerate(doc.paragraphs)
        if para_idx in allowed_toc_indexes
    }
    missing = [
        label
        for label in sorted(REQUIRED_TAIL_TOC_LABELS)
        if normalize(label) in page_map and normalize(label) not in present
    ]
    if missing:
        raise RuntimeError(
            "static TOC update cannot patch missing required tail entries; "
            "run canonical toc-tail-entries repair first: " + ", ".join(missing)
        )


def run_visible_text(run_element) -> str:
    return "".join(node.text or "" for node in run_element.findall(".//" + qn("w:t")))


def run_rpr_clone(run_element):
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        return None
    return deepcopy(rpr)


def run_font_attrs(run_element) -> dict[str, str]:
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        return {key: "" for key in FONT_ATTRS}
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return {key: "" for key in FONT_ATTRS}
    return {key: rfonts.attrib.get(qn(f"w:{key}"), "") for key in FONT_ATTRS}


def font_attrs_is_empty(attrs: dict[str, str]) -> bool:
    return not any(attrs.get(key) for key in FONT_ATTRS)


def iter_toc_text_runs_by_side(paragraph):
    seen_tab = False
    for run_element in iter_run_elements(paragraph):
        if run_element.find(qn("w:tab")) is not None:
            seen_tab = True
        text = run_visible_text(run_element)
        if not text.strip():
            continue
        yield ("post" if seen_tab else "pre"), run_element


def first_font_attrs_by_side(paragraph) -> dict[str, dict[str, str]]:
    result: dict[str, dict[str, str]] = {}
    for side, run_element in iter_toc_text_runs_by_side(paragraph):
        result.setdefault(side, run_font_attrs(run_element))
    return result


def first_rpr_by_side(paragraph) -> dict[str, object]:
    result: dict[str, object] = {}
    for side, run_element in iter_toc_text_runs_by_side(paragraph):
        result.setdefault(side, run_rpr_clone(run_element))
    return result


def iter_toc_runs_by_role(paragraph):
    seen_tab = False
    for run_element in iter_run_elements(paragraph):
        has_tab = run_element.find(qn("w:tab")) is not None
        text = run_visible_text(run_element)
        if has_tab:
            yield "tab", run_element
            seen_tab = True
        if text.strip():
            yield ("post" if seen_tab else "pre"), run_element


def first_rpr_by_role(paragraph) -> dict[str, object]:
    result: dict[str, object] = {}
    for role, run_element in iter_toc_runs_by_role(paragraph):
        result.setdefault(role, run_rpr_clone(run_element))
    return result


def toc_entry_level(paragraph) -> int | None:
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    style_id = (paragraph.style.style_id if paragraph.style else "").lower()
    text = paragraph.text.strip()
    label = text.split("\t", 1)[0].strip()
    page_match = re.search(r"(\d+)\s*$", text)
    if "\u2026" in label:
        label = label.split("\u2026", 1)[0].strip()
    if "\t" not in text and page_match:
        label = re.sub(r"\s*\d+\s*$", "", label).strip()
    level = heading_level(label or text)
    if level is None and normalize(label).lower() in {normalize("\u6458\u8981").lower(), normalize("Abstract").lower()}:
        level = 1
    style_key = f"{style_id} {style_name}"
    if level is None and ("toc" in style_key or "wpsoffice" in style_key or "\u76ee\u5f55" in style_key):
        digits = re.findall(r"\d+", style_key)
        level = int(digits[-1]) if digits else 1
    return level if level is not None else None


def toc_visible_label(text: str) -> str:
    label = str(text or "").split("\t", 1)[0].strip()
    if "\u2026" in label:
        label = label.split("\u2026", 1)[0].strip()
    return re.sub(r"\s*(?:\d+|[ivxlcdmIVXLCDM]+)\s*$", "", label).strip()


def is_front_matter_toc_label(text: str) -> bool:
    label = normalize(toc_visible_label(text)).lower()
    return label in {normalize("\u6458\u8981").lower(), normalize("Abstract").lower()}


def is_toc_leader_entry(text: str) -> bool:
    return "\u2026" in str(text or "") and re.search(r"\d+\s*$", str(text or "")) is not None


def is_body_logical_page_start_candidate(text: str) -> bool:
    """Return true for headings that start the body page-number sequence."""
    stripped = re.sub(r"^[\s\u25a1]+", "", str(text or "").strip())
    return bool(
        re.match(r"^第[0-9一二三四五六七八九十]+章", stripped)
        or re.match(r"^\d{1,2}[\s\u25a1]+\S", stripped)
    )


def collect_toc_font_baselines(reference_docx: Path) -> dict[int, dict[str, dict[str, str]]]:
    baselines: dict[int, dict[str, dict[str, str]]] = {}
    doc = Document(reference_docx)
    toc_started = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if is_toc_heading_text(text):
            toc_started = True
            continue
        if not toc_started:
            continue
        if is_front_matter_toc_label(text):
            continue
        has_tab = paragraph_has_run_tab(para)
        style_name = (para.style.name if para.style else "").lower()
        if not has_tab and not style_name.startswith("toc"):
            if heading_level(text) == 1:
                break
            continue
        level = toc_entry_level(para)
        if level is None:
            continue
        baselines.setdefault(level, first_font_attrs_by_side(para))
    return baselines


def collect_toc_rpr_baselines(reference_docx: Path) -> dict[int, dict[str, object]]:
    baselines: dict[int, dict[str, object]] = {}
    doc = Document(reference_docx)
    toc_started = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if is_toc_heading_text(text):
            toc_started = True
            continue
        if not toc_started:
            continue
        if is_front_matter_toc_label(text):
            continue
        has_tab = paragraph_has_run_tab(para)
        style_name = (para.style.name if para.style else "").lower()
        if not has_tab and not style_name.startswith("toc"):
            if heading_level(text) == 1:
                break
            continue
        level = toc_entry_level(para)
        if level is None:
            continue
        baselines.setdefault(level, first_rpr_by_role(para))
    return baselines


def apply_run_font_attrs(run_element, attrs: dict[str, str]) -> None:
    rpr = run_element.find(qn("w:rPr"))
    created_rpr = False
    if rpr is None:
        if font_attrs_is_empty(attrs):
            return
        rpr = OxmlElement("w:rPr")
        run_element.insert(0, rpr)
        created_rpr = True
    rfonts = rpr.find(qn("w:rFonts"))
    if font_attrs_is_empty(attrs):
        if rfonts is not None:
            rpr.remove(rfonts)
        if created_rpr and len(rpr) == 0:
            run_element.remove(rpr)
        return
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in FONT_ATTRS:
        attr = qn(f"w:{key}")
        value = attrs.get(key, "")
        if value:
            rfonts.set(attr, preferred_font(value))
        elif attr in rfonts.attrib:
            del rfonts.attrib[attr]


def apply_run_rpr(run_element, rpr_baseline) -> None:
    existing = run_element.find(qn("w:rPr"))
    if existing is not None:
        run_element.remove(existing)
    if rpr_baseline is not None:
        cloned = deepcopy(rpr_baseline)
        rfonts = cloned.find(qn("w:rFonts"))
        if rfonts is not None:
            for key in FONT_ATTRS:
                attr = qn(f"w:{key}")
                if attr in rfonts.attrib:
                    rfonts.set(attr, preferred_font(rfonts.attrib[attr]))
        run_element.insert(0, cloned)


def remove_web_hidden(run_element) -> None:
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        return
    web_hidden = rpr.find(qn("w:webHidden"))
    if web_hidden is not None:
        rpr.remove(web_hidden)
    if len(rpr) == 0:
        run_element.remove(rpr)


def unhide_static_toc_page_runs(paragraph) -> None:
    seen_tab = False
    for run_element in iter_run_elements(paragraph):
        has_tab = run_element.find(qn("w:tab")) is not None
        if has_tab:
            remove_web_hidden(run_element)
            seen_tab = True
            continue
        if not seen_tab:
            continue
        if run_element.find(qn("w:instrText")) is not None or run_element.find(qn("w:fldChar")) is not None:
            continue
        if run_visible_text(run_element).strip():
            remove_web_hidden(run_element)


def restore_toc_entry_run_fonts(paragraph, baselines: dict[int, dict[str, dict[str, str]]]) -> None:
    level = toc_entry_level(paragraph)
    if level is None or level not in baselines:
        return
    baseline = baselines[level]
    for side, run_element in iter_toc_text_runs_by_side(paragraph):
        attrs = baseline.get(side) or baseline.get("pre")
        if attrs is None:
            continue
        apply_run_font_attrs(run_element, attrs)


def restore_toc_entry_run_rpr(paragraph, baselines: dict[int, dict[str, object]]) -> None:
    level = toc_entry_level(paragraph)
    if level is None or level not in baselines:
        return
    baseline = baselines[level]
    for role, run_element in iter_toc_runs_by_role(paragraph):
        if role == "pre":
            rpr_baseline = baseline.get("pre")
        else:
            rpr_baseline = baseline.get(role)
        if rpr_baseline is None and role == "post":
            rpr_baseline = baseline.get("pre")
        if rpr_baseline is None and role == "tab":
            rpr_baseline = baseline.get("post") or baseline.get("pre")
        if rpr_baseline is None and role != "pre":
            rpr_baseline = baseline.get("pre") or baseline.get("post") or baseline.get("tab")
        apply_run_rpr(run_element, rpr_baseline)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--mapping", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument(
        "--reference-toc",
        required=True,
        help="DOCX whose TOC entry run-level fonts are the locked donor baseline.",
    )
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    mapping_path = Path(args.mapping).resolve()
    output_path = Path(args.output).resolve()
    reference_toc = Path(args.reference_toc).resolve()
    if not reference_toc.exists():
        raise FileNotFoundError(f"locked reference TOC donor does not exist: {reference_toc}")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with mapping_path.open("r", encoding="utf-8-sig") as f:
        rows = json.load(f)

    numeric_rows = []
    for row in rows:
        try:
            page = int(row["page"])
        except Exception:
            continue
        numeric_rows.append({**row, "page": page})
    body_start = min(
        (row["page"] for row in numeric_rows if is_body_logical_page_start_candidate(str(row.get("text") or ""))),
        default=None,
    )
    if body_start is None:
        body_start = min(
            (row["page"] for row in numeric_rows if int(row.get("level") or 1) == 1),
            default=None,
        )
    page_map = {}
    for row in numeric_rows:
        page = row["page"]
        logical_page = page - body_start + 1 if body_start is not None else page
        page_map[normalize(str(row["text"]))] = str(max(1, logical_page))

    doc = Document(input_path)
    toc_font_baselines = collect_toc_font_baselines(reference_toc) if reference_toc.exists() else {}
    toc_rpr_baselines = collect_toc_rpr_baselines(reference_toc) if reference_toc.exists() else {}
    allowed_toc_indexes = toc_paragraph_indexes(doc)
    require_mapped_tail_entries_present(doc, allowed_toc_indexes, page_map)
    for para_idx, para in enumerate(doc.paragraphs):
        if para_idx not in allowed_toc_indexes:
            continue
        style_name = (para.style.name if para.style else "").lower()
        has_tab = paragraph_has_run_tab(para)
        entry_text = toc_visible_label(para.text)
        leader_entry = is_toc_leader_entry(para.text)
        if style_name not in {"toc 1", "toc 2", "toc 3"} and not has_tab and not leader_entry:
            continue
        page = page_map.get(normalize(entry_text))
        if page:
            set_toc_page_preserving_runs(para, page)
        restore_toc_entry_run_fonts(para, toc_font_baselines)
        restore_toc_entry_run_rpr(para, toc_rpr_baselines)
        unhide_static_toc_page_runs(para)
    normalize_toc_right_tab_positions(doc, allowed_toc_indexes)
    normalize_toc_entry_paragraph_spacing(doc, allowed_toc_indexes)

    doc.save(output_path)
    keep_only_toc_document_xml_delta(input_path, output_path, allowed_toc_indexes)
    assert_only_allowed_package_drift(input_path, output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
