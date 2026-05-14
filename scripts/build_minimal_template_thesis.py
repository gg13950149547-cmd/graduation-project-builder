#!/usr/bin/env python3
"""Build a shortest manual-review thesis draft from a template without bypassing skill gates."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import time
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree as ET

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.document import Document as _Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

try:
    from win32com.client import constants, gencache  # type: ignore
except ImportError:  # pragma: no cover
    constants = None
    gencache = None

from docx_apply_table_family import patch_table
from docx_formula_number_table import patch_formula
from python_runtime import resolve_python_exe


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
PYTHON_EXE = resolve_python_exe()
HEADING_COLLECTION_TIMEOUT = 600
TOC_SYNC_TIMEOUT = 600
WORD_EXPORT_TIMEOUT = 600
MAX_FIGURE_WIDTH_CM = 14.2
MAX_FIGURE_HEIGHT_CM = 17.0
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
}
W = "{%s}" % NS["w"]
W14 = "{%s}" % NS["w14"]
ET.register_namespace("w", NS["w"])
ET.register_namespace("w14", NS["w14"])

CN_NUMBER_CHARS = "\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341"
BODY_HEADING_RE = re.compile(
    r"^(?:\d+(?:\.\d+){0,3}\s+\S.*|"
    r"\u7b2c[0-9\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+[\u7ae0\u8282]\s*\S*)$"
)
TABLE_TITLE_RE = re.compile(rf"^\s*(?:\u8868|\u7eed\u8868)\s*[0-9{CN_NUMBER_CHARS}]+(?:[-.]\d+)?[\s\u25a1]*\S")
FIGURE_TITLE_RE = re.compile(rf"^\s*\u56fe\s*[0-9{CN_NUMBER_CHARS}]+(?:[-.]\d+)?[\s\u25a1]*\S")
TAIL_HEADING_NORMALIZED = {
    "\u53c2\u8003\u6587\u732e",
    "\u81f4\u8c22",
    "\u9644\u5f55",
    "acknowledgements",
    "acknowledgments",
    "appendix",
}

CHINESE_ABSTRACT = "\u6458\u8981"
ENGLISH_ABSTRACT = "abstract"
CHINESE_KEYWORDS_LABEL = "\u5173\u952e\u8bcd\uff1a"
ENGLISH_KEYWORDS_LABEL = "Key Words: "
REFERENCES_HEADING = "\u53c2\u8003\u6587\u732e"
ACK_HEADING = "\u81f4\u8c22"
APPENDIX_HEADING = "\u9644\u5f55"
FORMULA_ANCHOR = "__GPB_MINIMAL_FORMULA_ANCHOR__"
DEFAULT_FLOWCHART_DRAWIO = """<mxfile host="app.diagrams.net" modified="2026-04-22T00:00:00.000Z" agent="Codex" version="24.7.17">
  <diagram id="page1" name="Page-1">
    <mxGraphModel dx="1200" dy="1200" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="850" pageHeight="1100" math="0" shadow="0">
      <root>
        <mxCell id="0"/>
        <mxCell id="1" parent="0"/>
        <mxCell id="2" value="开始" style="ellipse;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#000000;fontColor=#000000;fontSize=16;" vertex="1" parent="1">
          <mxGeometry x="300" y="60" width="160" height="70" as="geometry"/>
        </mxCell>
        <mxCell id="3" value="采集数据" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#000000;fontColor=#000000;fontSize=16;" vertex="1" parent="1">
          <mxGeometry x="270" y="190" width="220" height="80" as="geometry"/>
        </mxCell>
        <mxCell id="4" value="分析行为" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#000000;fontColor=#000000;fontSize=16;" vertex="1" parent="1">
          <mxGeometry x="270" y="340" width="220" height="80" as="geometry"/>
        </mxCell>
        <mxCell id="5" value="输出结果" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#000000;fontColor=#000000;fontSize=16;" vertex="1" parent="1">
          <mxGeometry x="270" y="490" width="220" height="80" as="geometry"/>
        </mxCell>
        <mxCell id="6" value="结束" style="ellipse;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#000000;fontColor=#000000;fontSize=16;" vertex="1" parent="1">
          <mxGeometry x="300" y="640" width="160" height="70" as="geometry"/>
        </mxCell>
        <mxCell id="8" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeColor=#000000;" edge="1" parent="1" source="2" target="3">
          <mxGeometry relative="1" as="geometry"/>
        </mxCell>
        <mxCell id="9" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeColor=#000000;" edge="1" parent="1" source="3" target="4">
          <mxGeometry relative="1" as="geometry"/>
        </mxCell>
        <mxCell id="10" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeColor=#000000;" edge="1" parent="1" source="4" target="5">
          <mxGeometry relative="1" as="geometry"/>
        </mxCell>
        <mxCell id="11" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeColor=#000000;" edge="1" parent="1" source="5" target="6">
          <mxGeometry relative="1" as="geometry"/>
        </mxCell>
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>
"""


@dataclass
class ParagraphDonor:
    ppr: ET.Element | None
    rpr: ET.Element | None
    style_id: str
    style_name: str


def run(cmd: list[str], *, timeout: int = 1200, check: bool = True) -> subprocess.CompletedProcess[str]:
    proc = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=timeout,
    )
    if check and proc.returncode != 0:
        raise RuntimeError(
            f"command failed: {' '.join(cmd)}\nstdout:\n{proc.stdout}\nstderr:\n{proc.stderr}"
        )
    return proc


def normalize_text(text: str) -> str:
    return re.sub(r"[\s\u25a1]+", "", text or "").strip().lower()


def heading_level_text(text: str) -> int | None:
    stripped = text.strip()
    if re.match(r"^\d{1,2}\s+\S", stripped):
        return 1
    if re.match(r"^\d{1,2}\.\d+\s+\S", stripped):
        return 2
    if re.match(r"^\d{1,2}\.\d+\.\d+\s+\S", stripped):
        return 3
    if re.match(r"^\d{1,2}\.\d+\.\d+\.\d+\s+\S", stripped):
        return 4
    if re.match(rf"^\u7b2c[0-9{CN_NUMBER_CHARS}]+[\u7ae0]\s*\S*", stripped):
        return 1
    return None


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        block_parent = parent._body
    else:
        parent_elm = parent._tc
        block_parent = parent
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, block_parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, block_parent)


def insert_paragraph_after(anchor: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def insert_paragraph_after_table(table: Table) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._element.addnext(new_p)
    return Paragraph(new_p, table._parent)


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def paragraph_index(doc: Document, target: Paragraph) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph._element is target._element:
            return idx
    raise KeyError("paragraph not found")


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def capture_donor(paragraph: Paragraph | None) -> ParagraphDonor:
    if paragraph is None:
        return ParagraphDonor(None, None, "", "")
    ppr = deepcopy(paragraph._element.pPr) if paragraph._element.pPr is not None else None
    rpr = None
    for run in paragraph.runs:
        if run._element.rPr is not None:
            rpr = deepcopy(run._element.rPr)
            break
    style_id = paragraph.style.style_id if paragraph.style is not None else ""
    style_name = paragraph.style.name if paragraph.style is not None else ""
    return ParagraphDonor(ppr, rpr, style_id, style_name)


def ensure_pstyle(paragraph: Paragraph, style_id: str) -> None:
    if not style_id:
        return
    ppr = paragraph._element.pPr
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    style_node = ppr.find(qn("w:pStyle"))
    if style_node is None:
        style_node = OxmlElement("w:pStyle")
        ppr.insert(0, style_node)
    style_node.set(qn("w:val"), style_id)


def clear_paragraph_structure_residue(paragraph: Paragraph) -> None:
    ppr = paragraph._element.pPr
    if ppr is None:
        return
    for tag in ("w:numPr", "w:outlineLvl"):
        node = ppr.find(qn(tag))
        if node is not None:
            ppr.remove(node)


def apply_donor(paragraph: Paragraph, donor: ParagraphDonor) -> None:
    existing_ppr = paragraph._element.pPr
    if existing_ppr is not None:
        paragraph._element.remove(existing_ppr)
    if donor.ppr is not None:
        paragraph._element.insert(0, deepcopy(donor.ppr))
    clear_paragraph_content(paragraph)


def usable_figure_width_cm_for_paragraph(paragraph: Paragraph) -> float:
    try:
        doc = paragraph.part.document
        section = doc.sections[0]
        width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    except Exception:
        width_cm = MAX_FIGURE_WIDTH_CM
    return max(4.0, min(MAX_FIGURE_WIDTH_CM, float(width_cm) - 0.2))


def clamp_image_width_cm(width_cm: float, image_path: Path, usable_width_cm: float) -> float:
    safe_width = min(float(width_cm), usable_width_cm)
    try:
        with Image.open(image_path) as image:
            if image.height > 0:
                width_by_height = MAX_FIGURE_HEIGHT_CM * image.width / image.height
                safe_width = min(safe_width, width_by_height)
    except Exception:
        pass
    return max(4.0, safe_width)


def safe_image_width_cm(width_cm: float, image_path: Path, anchor: Paragraph) -> float:
    return clamp_image_width_cm(width_cm, image_path, usable_figure_width_cm_for_paragraph(anchor))


def add_run_with_donor(paragraph: Paragraph, donor: ParagraphDonor, text: str):
    run_obj = paragraph.add_run(text)
    existing_rpr = run_obj._element.rPr
    if existing_rpr is not None:
        run_obj._element.remove(existing_rpr)
    if donor.rpr is not None:
        run_obj._element.insert(0, deepcopy(donor.rpr))
    return run_obj


def set_run_font(run_obj, *, east_asia: str, size_pt: float, bold: bool = False) -> None:
    run_obj.font.name = "Times New Roman"
    run_obj.font.size = Pt(size_pt)
    run_obj.bold = bold
    rpr = run_obj._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), "Times New Roman")


def replace_single_run_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run_obj in paragraph.runs[1:]:
            run_obj.text = ""
    else:
        paragraph.add_run(text)


def set_keyword_line(paragraph: Paragraph, label: str, content: str) -> None:
    while len(paragraph.runs) < 2:
        paragraph.add_run("")
    paragraph.runs[0].text = label
    paragraph.runs[1].text = content
    for run_obj in paragraph.runs[2:]:
        run_obj.text = ""


def apply_run_font_pattern(paragraph: Paragraph, specs: list[tuple[int, str, float, bool]]) -> None:
    for index, east_asia, size_pt, bold in specs:
        while len(paragraph.runs) <= index:
            paragraph.add_run("")
        run_obj = paragraph.runs[index]
        set_run_font(run_obj, east_asia=east_asia, size_pt=size_pt, bold=bold)


def is_heading1(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    text = paragraph.text.strip()
    if not text:
        return False
    if "heading 1" in style_name or "\u6807\u9898 1" in style_name:
        return True
    compact = normalize_text(text)
    if compact in TAIL_HEADING_NORMALIZED:
        return True
    return bool(BODY_HEADING_RE.match(text))


def is_any_heading(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    text = paragraph.text.strip()
    if not text:
        return False
    if "heading" in style_name or "\u6807\u9898" in style_name:
        return True
    return bool(BODY_HEADING_RE.match(text))


def is_tail_heading(paragraph: Paragraph) -> bool:
    return normalize_text(paragraph.text) in TAIL_HEADING_NORMALIZED


def is_toc_paragraph(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    return style_name.startswith("toc")


def collect_body_headings(doc: Document) -> list[Paragraph]:
    return [
        paragraph
        for paragraph in doc.paragraphs
        if "\t" not in paragraph.text
        and is_heading1(paragraph)
        and not is_tail_heading(paragraph)
        and not is_toc_paragraph(paragraph)
    ]


def find_heading(doc: Document, normalized_heading: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if normalize_text(paragraph.text) == normalized_heading:
            return paragraph
    return None


def next_nonempty_paragraph(doc: Document, anchor: Paragraph) -> Paragraph | None:
    start = paragraph_index(doc, anchor)
    for paragraph in doc.paragraphs[start + 1 :]:
        if paragraph.text.strip():
            return paragraph
    return None


def first_nonheading_body_paragraph(doc: Document, heading: Paragraph) -> Paragraph | None:
    start = paragraph_index(doc, heading)
    for paragraph in doc.paragraphs[start + 1 :]:
        text = paragraph.text.strip()
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if text and paragraph is not heading and (
            "heading 1" in style_name or "\u6807\u9898 1" in style_name or heading_level_text(text) == 1
        ):
            break
        if text and not is_any_heading(paragraph):
            return paragraph
    return None


def find_caption_donor(doc: Document) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if FIGURE_TITLE_RE.match(text):
            return paragraph
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if TABLE_TITLE_RE.match(text):
            return paragraph
        if any(token in style_name for token in ("caption", "\u56fe\u9898", "\u8868\u9898")):
            return paragraph
    return None


def find_image_holder_donor(doc: Document, caption_paragraph: Paragraph | None) -> Paragraph | None:
    if caption_paragraph is None:
        return None
    anchor_index = paragraph_index(doc, caption_paragraph)
    for paragraph in reversed(doc.paragraphs[:anchor_index]):
        has_picture = any(run._element.xpath(".//w:drawing") for run in paragraph.runs)
        if has_picture:
            return paragraph
        if paragraph.text.strip():
            break
    return None


def strip_body_to_headings(doc: Document) -> None:
    body_headings = collect_body_headings(doc)
    if not body_headings:
        raise RuntimeError("template does not expose any body chapter headings")
    first_heading = body_headings[0]
    preserved = {
        paragraph._element
        for paragraph in doc.paragraphs
        if is_heading1(paragraph) and not is_toc_paragraph(paragraph)
    }
    started = False
    body = doc.element.body
    for child in list(body.iterchildren()):
        if child is first_heading._element:
            started = True
            continue
        if not started:
            continue
        if child.tag == W + "sectPr":
            continue
        if child.tag == W + "p" and child in preserved:
            continue
        body.remove(child)


def prepare_body_paragraph(paragraph: Paragraph, donor: ParagraphDonor) -> None:
    apply_donor(paragraph, donor)
    if donor.style_name:
        try:
            paragraph.style = donor.style_name
        except KeyError:
            pass
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    ensure_pstyle(paragraph, donor.style_id)


def prepare_center_paragraph(paragraph: Paragraph, donor: ParagraphDonor) -> None:
    apply_donor(paragraph, donor)
    if donor.style_name:
        try:
            paragraph.style = donor.style_name
        except KeyError:
            pass
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    ensure_pstyle(paragraph, donor.style_id)


def prepare_image_holder_paragraph(paragraph: Paragraph, donor: ParagraphDonor) -> None:
    apply_donor(paragraph, donor)
    if donor.style_name:
        try:
            paragraph.style = donor.style_name
        except KeyError:
            pass
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_with_next = True
    ensure_pstyle(paragraph, donor.style_id)
    clear_paragraph_structure_residue(paragraph)


def append_body_paragraph(anchor: Paragraph, donor: ParagraphDonor, text: str) -> Paragraph:
    paragraph = insert_paragraph_after(anchor)
    prepare_body_paragraph(paragraph, donor)
    add_run_with_donor(paragraph, donor, text)
    return paragraph


def append_reference_paragraph(anchor: Paragraph, donor: ParagraphDonor, text: str) -> Paragraph:
    paragraph = insert_paragraph_after(anchor)
    apply_donor(paragraph, donor)
    if donor.style_name:
        try:
            paragraph.style = donor.style_name
        except KeyError:
            pass
    ensure_pstyle(paragraph, donor.style_id)
    add_run_with_donor(paragraph, donor, text)
    return paragraph


def append_center_title(anchor: Paragraph, donor: ParagraphDonor, text: str) -> Paragraph:
    paragraph = insert_paragraph_after(anchor)
    prepare_center_paragraph(paragraph, donor)
    clear_paragraph_structure_residue(paragraph)
    add_run_with_donor(paragraph, donor, text)
    return paragraph


def append_image_block(
    anchor: Paragraph,
    holder_donor: ParagraphDonor,
    caption_donor: ParagraphDonor,
    image_path: Path,
    caption_text: str,
    *,
    width_cm: float,
) -> Paragraph:
    holder = insert_paragraph_after(anchor)
    prepare_image_holder_paragraph(holder, holder_donor)
    holder.add_run().add_picture(str(image_path), width=Cm(safe_image_width_cm(width_cm, image_path, anchor)))

    caption = insert_paragraph_after(holder)
    prepare_center_paragraph(caption, caption_donor)
    add_run_with_donor(caption, caption_donor, caption_text)
    caption.paragraph_format.space_after = Pt(0)
    return caption


def append_svg_placeholder(
    anchor: Paragraph,
    holder_donor: ParagraphDonor,
    caption_donor: ParagraphDonor,
    caption_text: str,
) -> tuple[Paragraph, Paragraph]:
    holder = insert_paragraph_after(anchor)
    prepare_image_holder_paragraph(holder, holder_donor)

    caption = insert_paragraph_after(holder)
    prepare_center_paragraph(caption, caption_donor)
    add_run_with_donor(caption, caption_donor, caption_text)
    caption.paragraph_format.space_after = Pt(0)
    return holder, caption


def append_table(
    document: Document,
    anchor: Paragraph,
    title_donor: ParagraphDonor,
    data: list[list[str]],
    title_text: str,
) -> Paragraph:
    title_paragraph = append_center_title(anchor, title_donor, title_text)
    title_paragraph.paragraph_format.keep_with_next = True
    clear_paragraph_structure_residue(title_paragraph)
    table = document.add_table(rows=len(data), cols=len(data[0]))
    title_paragraph._element.addnext(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(data):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            paragraph = cell.paragraphs[0]
            clear_paragraph_content(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else (WD_ALIGN_PARAGRAPH.LEFT if col_index == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run_obj = paragraph.add_run(value)
            set_run_font(run_obj, east_asia="SimSun", size_pt=10.5, bold=row_index == 0)
    return insert_paragraph_after_table(table)


def insert_formula_omml(docx_path: Path, anchor_text: str, para_id: str) -> None:
    with zipfile.ZipFile(docx_path) as zin:
        members = {name: zin.read(name) for name in zin.namelist()}
    root = ET.fromstring(members["word/document.xml"])
    body = root.find("w:body", NS)
    if body is None:
        raise RuntimeError("word/document.xml missing body")
    children = list(body)
    anchor_index = None
    for idx, child in enumerate(children):
        if child.tag != W + "p":
            continue
        text = "".join((node.text or "") for node in child.findall(".//w:t", NS)).strip()
        if text == anchor_text:
            anchor_index = idx
            break
    if anchor_index is None:
        raise RuntimeError(f"formula anchor not found: {anchor_text}")
    formula_xml = (
        f'<w:p xmlns:w="{NS["w"]}" xmlns:w14="{NS["w14"]}" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        f'w14:paraId="{para_id}" w14:textId="00ABCDEF"><m:oMathPara><m:oMath><m:r><m:t>y=\\beta_0+\\beta_1x</m:t></m:r></m:oMath></m:oMathPara></w:p>'
    )
    formula_paragraph = ET.fromstring(formula_xml)
    body.insert(anchor_index, formula_paragraph)
    body.remove(children[anchor_index])
    members["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in members.items():
            zout.writestr(name, data)


def bind_body_citation_paragraphs_to_style(docx_path: Path, style_id: str) -> None:
    if not style_id:
        return
    with zipfile.ZipFile(docx_path) as zin:
        members = {name: zin.read(name) for name in zin.namelist()}
    root = ET.fromstring(members["word/document.xml"])
    body = root.find("w:body", NS)
    if body is None:
        return

    body_started = False
    for paragraph in body.findall("w:p", NS):
        text = "".join((node.text or "") for node in paragraph.findall(".//w:t", NS)).strip()
        if not text:
            continue
        if BODY_HEADING_RE.match(text):
            body_started = True
            continue
        if normalize_text(text) in {normalize_text(REFERENCES_HEADING), normalize_text("references")}:
            break
        if not body_started:
            continue
        if "[" not in text or "]" not in text:
            continue
        if FIGURE_TITLE_RE.match(text) or TABLE_TITLE_RE.match(text):
            continue
        ppr = paragraph.find(W + "pPr")
        if ppr is None:
            ppr = ET.Element(W + "pPr")
            paragraph.insert(0, ppr)
        style_node = ppr.find(W + "pStyle")
        if style_node is None:
            style_node = ET.Element(W + "pStyle")
            ppr.insert(0, style_node)
        style_node.set(W + "val", style_id)

    members["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in members.items():
            zout.writestr(name, data)


def is_formula_layout_table(table: ET.Element) -> bool:
    return any("officeDocument/2006/math" in node.tag for node in table.iter())


def rendered_table_indices(docx_path: Path) -> list[int]:
    with zipfile.ZipFile(docx_path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    body = root.find("w:body", NS)
    if body is None:
        return []
    table_indices: list[int] = []
    total_tables = 0
    previous_text = ""
    for child in list(body):
        if child.tag == W + "p":
            previous_text = "".join((node.text or "") for node in child.findall(".//w:t", NS)).strip()
        elif child.tag == W + "tbl":
            if is_formula_layout_table(child):
                continue
            total_tables += 1
            if TABLE_TITLE_RE.match(previous_text.replace("\xa0", " ")):
                table_indices.append(total_tables)
    return table_indices or list(range(1, total_tables + 1))


def load_font(size: int):
    for candidate in (
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttf"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ):
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def make_dashboard(path: Path, title: str, accent: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1280, 720), "#F3F4F6")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34)
    text_font = load_font(22)
    value_font = load_font(36)
    small_font = load_font(18)
    draw.rounded_rectangle((36, 30, 1244, 116), radius=22, fill=accent)
    draw.text((64, 54), title, fill="white", font=title_font)
    cards = [
        (72, 160, 372, 328, "Active Videos", "12"),
        (404, 160, 704, 328, "Processed Tasks", "24"),
        (736, 160, 1036, 328, "Focus Score", "88%"),
        (1068, 160, 1208, 328, "Alerts", "3"),
    ]
    for x1, y1, x2, y2, label, value in cards:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill="white", outline="#CBD5E1", width=3)
        draw.text((x1 + 20, y1 + 18), label, fill="#475569", font=text_font)
        draw.text((x1 + 20, y1 + 92), value, fill="#0F172A", font=value_font)
    panels = [
        (72, 372, 612, 658, "Trend View"),
        (668, 372, 1208, 658, "Overview"),
    ]
    for x1, y1, x2, y2, label in panels:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill="white", outline="#CBD5E1", width=3)
        draw.text((x1 + 20, y1 + 16), label, fill="#334155", font=text_font)
        draw.line((x1 + 40, y2 - 72, x1 + 150, y1 + 120), fill="#2563EB", width=5)
        draw.line((x1 + 150, y1 + 120, x1 + 272, y2 - 150), fill="#2563EB", width=5)
        draw.line((x1 + 272, y2 - 150, x1 + 392, y1 + 170), fill="#2563EB", width=5)
        draw.text((x1 + 20, y2 - 36), "placeholder screenshot for template smoke review", fill="#64748B", font=small_font)
    image.save(path)


def export_drawio(drawio_path: Path, svg_path: Path, png_path: Path) -> None:
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "export_thesis_drawio_figure.py"),
            "--input-drawio",
            str(drawio_path),
            "--output-svg",
            str(svg_path),
            "--output-png",
            str(png_path),
            "--width",
            "900",
        ],
        timeout=600,
    )


def resolve_svg_requests(doc: Document, requests: list[dict[str, object]]) -> list[dict[str, object]]:
    resolved: list[dict[str, object]] = []
    for request in requests:
        paragraph = request["paragraph"]
        if not isinstance(paragraph, Paragraph):
            continue
        resolved.append(
            {
                "paragraph_index": paragraph_index(doc, paragraph) + 1,
                "svg_path": Path(str(request["svg_path"])),
                "fallback_png": Path(str(request["fallback_png"])),
                "width_cm": float(request["width_cm"]),
            }
        )
    return resolved


def svg_height_cm(fallback_png: Path, width_cm: float) -> float:
    with Image.open(fallback_png) as image:
        return width_cm * image.height / image.width


def apply_svg_pictures(
    docx_path: Path,
    requests: list[dict[str, object]],
    *,
    holder_donor: ParagraphDonor | None = None,
    caption_donor: ParagraphDonor | None = None,
) -> None:
    for request in requests:
        fallback_png = Path(str(request["fallback_png"]))
        width_cm = clamp_image_width_cm(float(request["width_cm"]), fallback_png, MAX_FIGURE_WIDTH_CM)
        command = [
            "officecli",
            "add",
            str(docx_path),
            f"/body/p[{int(request['paragraph_index'])}]",
            "--type",
            "picture",
            "--prop",
            f"src={Path(str(request['svg_path']))}",
            "--prop",
            f"fallback={fallback_png}",
            "--prop",
            f"width={width_cm:.2f}cm",
            "--prop",
            f"height={svg_height_cm(fallback_png, width_cm):.2f}cm",
        ]
        run(command, timeout=600)
        subprocess.run(
            ["officecli", "close", str(docx_path)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=False,
            timeout=120,
        )
    wait_for_docx_ready(docx_path)
    if holder_donor is None or caption_donor is None:
        return
    doc = Document(str(docx_path))
    for request in requests:
        paragraph_idx = int(request["paragraph_index"]) - 1
        holder = doc.paragraphs[paragraph_idx]
        prepare_image_holder_paragraph(holder, holder_donor)
        if paragraph_idx + 1 < len(doc.paragraphs):
            caption = doc.paragraphs[paragraph_idx + 1]
            prepare_center_paragraph(caption, caption_donor)
            caption.paragraph_format.space_after = Pt(6)
    doc.save(str(docx_path))
    wait_for_docx_ready(docx_path)


def media_count(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path) as zf:
        return sum(1 for name in zf.namelist() if name.startswith("word/media/"))


def wait_for_docx_ready(docx_path: Path, *, retries: int = 30, delay_s: float = 0.5) -> None:
    last_error: Exception | None = None
    for _ in range(retries):
        try:
            with zipfile.ZipFile(docx_path) as zf:
                _ = zf.namelist()
            return
        except (PermissionError, OSError, zipfile.BadZipFile) as exc:
            last_error = exc
            time.sleep(delay_s)
    if last_error is not None:
        raise last_error


def sync_static_toc_if_possible(
    working_docx: Path,
    output_docx: Path,
    run_root: Path,
    reference_docx: Path | None = None,
) -> None:
    heading_pages = run_root / "meta" / "heading-pages.json"
    heading_pages.parent.mkdir(parents=True, exist_ok=True)
    collected = run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "collect_heading_pages_word.py"),
            "--input",
            str(working_docx),
            "--output",
            str(heading_pages),
        ],
        check=False,
        timeout=HEADING_COLLECTION_TIMEOUT,
    )
    if collected.returncode != 0 or not heading_pages.exists():
        raise RuntimeError("static TOC page sync failed: heading page collection did not produce evidence")

    media_before = media_count(working_docx) if working_docx.exists() else 0
    synced = run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "update_static_toc.py"),
            "--input",
            str(working_docx),
            "--mapping",
            str(heading_pages),
            "--output",
            str(output_docx),
            *(["--reference-toc", str(reference_docx)] if reference_docx is not None else []),
        ],
        check=False,
        timeout=TOC_SYNC_TIMEOUT,
    )
    media_after = media_count(output_docx) if output_docx.exists() else -1
    if synced.returncode != 0 or not output_docx.exists() or (media_before > 0 and media_after < media_before):
        raise RuntimeError("static TOC page sync failed: update_static_toc did not produce a safe output DOCX")


def export_with_word_direct(docx_path: Path, pdf_path: Path) -> None:
    if gencache is None or constants is None:
        raise RuntimeError("pywin32 is not available")
    word = None
    document = None
    try:
        word = gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(docx_path))
        document.Repaginate()
        try:
            document.Fields.Update()
        except Exception:
            pass
        try:
            for toc in document.TablesOfContents:
                toc.Update()
        except Exception:
            pass
        document.Save()
        document.ExportAsFixedFormat(str(pdf_path), constants.wdExportFormatPDF)
    finally:
        if document is not None:
            try:
                document.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def export_with_word(docx_path: Path, pdf_path: Path, *, timeout: int = WORD_EXPORT_TIMEOUT) -> None:
    run(
        [
            str(PYTHON_EXE),
            str(Path(__file__).resolve()),
            "--word-export-only",
            "--word-export-docx",
            str(docx_path),
            "--word-export-pdf",
            str(pdf_path),
        ],
        timeout=timeout,
    )


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        export_with_word(docx_path, pdf_path)
    except Exception:
        run(
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(SCRIPT_DIR / "wps_export_pdf.ps1"),
                "-InputDocx",
                str(docx_path),
                "-OutputPdf",
                str(pdf_path),
            ],
            timeout=600,
        )
    if not pdf_path.exists():
        raise RuntimeError(f"renderer did not create pdf: {pdf_path}")


def build_assets(run_root: Path) -> tuple[dict[str, str], Path]:
    assets_root = run_root / "assets"
    drawio_root = assets_root / "drawio"
    screenshot_root = assets_root / "screenshots"
    drawio_root.mkdir(parents=True, exist_ok=True)
    screenshot_root.mkdir(parents=True, exist_ok=True)

    drawio_path = drawio_root / "minimal-structure.drawio"
    svg_path = drawio_root / "minimal-structure.svg"
    png_path = drawio_root / "minimal-structure.png"
    drawio_path.write_text(DEFAULT_FLOWCHART_DRAWIO, encoding="utf-8")
    export_drawio(drawio_path, svg_path, png_path)

    screenshot_path = screenshot_root / "minimal-dashboard.png"
    make_dashboard(screenshot_path, "Minimal Thesis Smoke Dashboard", "#2563EB")

    asset_manifest = run_root / "meta" / "asset_manifest.json"
    asset_manifest.parent.mkdir(parents=True, exist_ok=True)
    asset_manifest.write_text(
        json.dumps(
            {
                "diagrams": {
                    "minimal_structure": {
                        "drawio": str(drawio_path),
                        "svg": str(svg_path),
                        "png": str(png_path),
                        "family": "flowchart",
                        "sample_lock": "figure-flowchart-vertical-sample-01",
                        "svg_primary_expected": True,
                    }
                },
                "screenshots": {
                    "minimal_dashboard": str(screenshot_path),
                },
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return {
        "diagram_png": str(png_path),
        "diagram_svg": str(svg_path),
        "screenshot_png": str(screenshot_path),
    }, asset_manifest


def build_minimal_doc(template_path: Path, working_docx: Path, run_root: Path) -> tuple[Path, Path]:
    assets, asset_manifest = build_assets(run_root)
    shutil.copy2(template_path, working_docx)
    doc = Document(str(working_docx))
    svg_requests: list[dict[str, object]] = []

    chinese_abstract = find_heading(doc, normalize_text(CHINESE_ABSTRACT))
    english_abstract = find_heading(doc, normalize_text(ENGLISH_ABSTRACT))
    references_heading = find_heading(doc, normalize_text(REFERENCES_HEADING))
    acknowledgement_heading = find_heading(doc, normalize_text(ACK_HEADING))
    appendix_heading = find_heading(doc, normalize_text(APPENDIX_HEADING))
    body_headings = collect_body_headings(doc)

    if chinese_abstract is None or english_abstract is None or references_heading is None:
        raise RuntimeError("template is missing required abstract or references headings")
    if not body_headings:
        raise RuntimeError("template is missing body headings")

    chinese_abstract_body = next_nonempty_paragraph(doc, chinese_abstract)
    english_abstract_body = next_nonempty_paragraph(doc, english_abstract)
    chinese_keywords = next_nonempty_paragraph(doc, chinese_abstract_body or chinese_abstract)
    english_keywords = next_nonempty_paragraph(doc, english_abstract_body or english_abstract)
    body_donor_paragraph = first_nonheading_body_paragraph(doc, body_headings[0])
    caption_donor_paragraph = find_caption_donor(doc)
    if caption_donor_paragraph is None:
        raise RuntimeError("template does not expose an approved caption donor paragraph")
    holder_donor_paragraph = find_image_holder_donor(doc, caption_donor_paragraph)
    if holder_donor_paragraph is None:
        raise RuntimeError("template does not expose an approved image-holder donor paragraph")
    reference_donor_paragraph = next_nonempty_paragraph(doc, references_heading)
    acknowledgement_donor_paragraph = next_nonempty_paragraph(doc, acknowledgement_heading) if acknowledgement_heading else None
    appendix_donor_paragraph = next_nonempty_paragraph(doc, appendix_heading) if appendix_heading else None

    body_donor = capture_donor(body_donor_paragraph)
    holder_donor = capture_donor(holder_donor_paragraph)
    caption_donor = capture_donor(caption_donor_paragraph)
    reference_donor = capture_donor(reference_donor_paragraph)
    acknowledgement_donor = capture_donor(acknowledgement_donor_paragraph)
    appendix_donor = capture_donor(appendix_donor_paragraph)

    strip_body_to_headings(doc)

    if chinese_abstract_body is not None:
        replace_single_run_text(
            chinese_abstract_body,
            "\u672c\u6587\u57fa\u4e8e\u7ed9\u5b9a\u6a21\u677f\u751f\u6210\u4e00\u7bc7\u6700\u77ed\u8bba\u6587\u6837\u7a3f\uff0c\u53ea\u4fdd\u7559\u6458\u8981\uff0c\u5404\u7ae0\u8282\uff0c\u7ed3\u6784\u56fe\uff0c\u622a\u56fe\uff0c\u4e09\u7ebf\u8868\uff0c\u516c\u5f0f\u548c\u53c2\u8003\u6587\u732e\u7b49\u5fc5\u8981\u6784\u4ef6\uff0c\u7528\u4e8e\u68c0\u67e5 skill \u5728\u6700\u5c0f\u5185\u5bb9\u573a\u666f\u4e0b\u7684\u751f\u6210\u8def\u5f84\u662f\u5426\u7a33\u5b9a\u3002",
        )
        apply_run_font_pattern(chinese_abstract_body, [(0, "SimSun", 12, False)])
    if chinese_keywords is not None:
        set_keyword_line(chinese_keywords, CHINESE_KEYWORDS_LABEL, "\u6a21\u677f\u5bf9\u9f50\uff1b\u6700\u77ed\u8bba\u6587\uff1b\u683c\u5f0f\u95e8\u7981")
        apply_run_font_pattern(chinese_keywords, [(0, "SimSun", 12, True), (1, "SimSun", 12, False)])
    if english_abstract_body is not None:
        replace_single_run_text(
            english_abstract_body,
            "This thesis is a shortest template-aligned smoke manuscript. It keeps the abstract surfaces, chapter skeleton, structural figure route, runtime screenshot placeholder, formula sample, table sample, and bibliography so the skill can be reviewed on a real template path.",
        )
        apply_run_font_pattern(english_abstract_body, [(0, "Times New Roman", 12, False)])
    if english_keywords is not None:
        set_keyword_line(english_keywords, ENGLISH_KEYWORDS_LABEL, "Template Alignment; Minimal Thesis; Gate Verification")
        apply_run_font_pattern(english_keywords, [(0, "Times New Roman", 12, True), (1, "Times New Roman", 12, False)])

    chapter_texts = [
        "\u672c\u7ae0\u4ea4\u4ee3\u6700\u77ed\u8bba\u6587 smoke \u6837\u7a3f\u7684\u76ee\u6807\uff0c\u4e3b\u8981\u7528\u4e8e\u68c0\u67e5\u6a21\u677f\u3001\u5f15\u7528\u548c\u540e\u7eed\u95e8\u7981\u94fe\u8def\u5728\u5c0f\u5185\u5bb9\u573a\u666f\u4e0b\u662f\u5426\u4f1a\u5931\u5b88[1]\u3002",
        "\u672c\u7ae0\u4fdd\u7559\u6700\u5c0f\u7406\u8bba\u6bb5\u843d\uff0c\u7528\u4e8e\u8bf4\u660e\u6a21\u677f\u5bf9\u9f50\u4e0d\u53ea\u662f\u9875\u9762\u5916\u89c2\uff0c\u8fd8\u9700\u8981\u8ba9\u6b63\u6587\u5f15\u7528\u3001\u6bb5\u843d\u548c\u5404\u7c7b\u975e\u6b63\u6587\u9762\u76f8\u4e92\u72ec\u7acb[2]\u3002",
        "\u672c\u7ae0\u63d2\u5165 draw.io \u751f\u6210\u7684\u7ed3\u6784\u56fe\uff0c\u7528\u4e8e\u9a8c\u8bc1\u7ed3\u6784\u56fe\u5fc5\u987b\u4fdd\u7559 draw.io \u6e90\u6587\u4ef6\u3001\u5bfc\u51fa\u8d44\u4ea7\u548c\u63d2\u5165\u540e\u7684\u9875\u9762\u53ef\u8bfb\u6027[3]\u3002",
        "\u672c\u7ae0\u4f7f\u7528\u8fd0\u884c\u754c\u9762\u622a\u56fe\u5360\u4f4d\u56fe\u4e0e\u4e09\u7ebf\u8868\u6837\u4f8b\uff0c\u68c0\u67e5\u622a\u56fe\u5bb6\u65cf\u548c\u8868\u683c\u5bb6\u65cf\u4e0d\u518d\u9519\u8bef\u7ee7\u627f\u6b63\u6587\u7f29\u8fdb\u6216\u6b63\u6587\u5b57\u4f53[4]\u3002",
        "\u672c\u7ae0\u63d2\u5165\u516c\u5f0f\u5bf9\u8c61\u548c\u53f3\u4fa7\u7f16\u53f7\uff0c\u7528\u4e8e\u9a8c\u8bc1\u516c\u5f0f\u884c\u72ec\u7acb\u4e8e\u6b63\u6587\u6bb5\u843d\uff0c\u4e14\u7f16\u53f7\u4f4d\u7f6e\u4fdd\u6301\u4e3a\u540c\u884c\u53f3\u4fa7\u9762[5]\u3002",
        "\u7ed3\u8bba\u90e8\u5206\u8bf4\u660e\uff0c\u6700\u77ed\u6837\u7a3f\u4e5f\u5fc5\u987b\u8d70\u5b8c\u5f15\u7528\u5f52\u4e00\u5316\uff0c\u81ea\u68c0\uff0c\u5b57\u4f53\u5ba1\u8ba1\u548c acceptance record \u751f\u6210\uff0c\u4e0d\u80fd\u518d\u51fa\u73b0\u53ea\u770b\u4e00\u773c\u683c\u5f0f\u6ca1\u95ee\u9898\u4f46\u6b63\u6587\u5df2\u7ecf\u6df7\u4e71\u7684\u60c5\u51b5[6]\u3002",
    ]

    formula_chapter_index = min(4, len(body_headings) - 1)
    figure_chapter_index = min(2, len(body_headings) - 1)
    screenshot_table_index = min(3, len(body_headings) - 1)

    for idx, heading in enumerate(body_headings):
        anchor = heading
        text = chapter_texts[min(idx, len(chapter_texts) - 1)]
        anchor = append_body_paragraph(anchor, body_donor, text)
        if idx == figure_chapter_index:
            anchor = append_image_block(
                anchor,
                holder_donor,
                caption_donor,
                Path(assets["diagram_png"]),
                "\u56fe 1 \u7cfb\u7edf\u5904\u7406\u6d41\u7a0b\u56fe",
                width_cm=11.5,
            )
            anchor = append_body_paragraph(
                anchor,
                body_donor,
                "\u8be5\u7ed3\u6784\u56fe\u5c55\u793a\u4e86\u6700\u77ed\u6837\u7a3f\u4e2d\u7ed3\u6784\u56fe\u3001\u5916\u90e8\u56fe\u9898\u4e0e\u56fe\u540e\u6b63\u6587\u4ecb\u7ecd\u5fc5\u987b\u4fdd\u6301\u540c\u4e00\u672c\u5730\u56fe\u5757\u7684\u57fa\u672c\u7ea6\u675f\u3002",
            )
        if idx == screenshot_table_index:
            anchor = append_image_block(
                anchor,
                holder_donor,
                caption_donor,
                Path(assets["screenshot_png"]),
                "\u56fe 2 \u7cfb\u7edf\u754c\u9762\u5360\u4f4d\u622a\u56fe",
                width_cm=11.0,
            )
            anchor = append_body_paragraph(
                anchor,
                body_donor,
                "\u8be5\u622a\u56fe\u5bf9\u5e94\u7684\u6b63\u6587\u4ecb\u7ecd\u7528\u4e8e\u9a8c\u8bc1\u622a\u56fe\u3001\u56fe\u9898\u548c\u540e\u7eed\u8bf4\u660e\u6bb5\u4e0d\u4f1a\u88ab\u5168\u5c40\u683c\u5f0f\u5316\u8def\u5f84\u6253\u6563\u3002",
            )
            anchor = append_table(
                doc,
                anchor,
                caption_donor,
                [
                    ["\u5b57\u6bb5", "\u542b\u4e49", "\u6765\u6e90"],
                    ["video_name", "\u89c6\u9891\u540d\u79f0", "\u4e0a\u4f20\u8bb0\u5f55"],
                    ["task_state", "\u4efb\u52a1\u72b6\u6001", "\u540e\u53f0\u5904\u7406"],
                    ["focus_score", "\u4e13\u6ce8\u5ea6\u5f97\u5206", "\u8bc4\u4f30\u7ed3\u679c"],
                ],
                "\u8868 1 \u6700\u5c0f\u6837\u4f8b\u6570\u636e\u8868",
            )
        if idx == formula_chapter_index:
            formula_anchor = append_center_title(anchor, caption_donor, FORMULA_ANCHOR)
            formula_anchor.paragraph_format.space_before = Pt(6)
            formula_anchor.paragraph_format.space_after = Pt(6)

    if references_heading is not None:
        anchor = references_heading
        bibliography_entries = [
            "[1] Example Research on Template-Aligned Thesis Delivery[J].",
            "[2] Sample Study on Rendered-Page Review[J].",
            "[3] Minimal Thesis Diagram Governance Practice[J].",
            "[4] Screenshot and Table Surface Isolation Study[J].",
            "[5] Formula Numbering Surface Audit Example[J].",
            "[6] Document Gate Regression Checklist[J].",
        ]
        for line in bibliography_entries:
            anchor = append_reference_paragraph(anchor, reference_donor, line)
    if acknowledgement_heading is not None:
        append_body_paragraph(
            acknowledgement_heading,
            acknowledgement_donor if acknowledgement_donor.ppr is not None else body_donor,
            "\u611f\u8c22\u672c\u8f6e skill audit \u4e2d\u63d0\u4f9b\u771f\u5b9e\u6a21\u677f\u3001\u5b9e\u9645\u9519\u8bef\u56de\u9988\u548c\u4eba\u5de5\u590d\u6838\u610f\u89c1\u7684\u4e0a\u4e0b\u6587\u3002",
        )
    if appendix_heading is not None:
        append_body_paragraph(
            appendix_heading,
            appendix_donor if appendix_donor.ppr is not None else body_donor,
            "\u9644\u5f55\u4fdd\u7559\u4e3a\u6700\u5c0f smoke \u6837\u7a3f\u8bf4\u660e\uff0c\u540e\u7eed\u53ef\u7ee7\u7eed\u586b\u5145\u811a\u672c\u3001\u6570\u636e\u5b57\u6bb5\u6216\u8bc4\u5ba1\u8bb0\u5f55\u3002",
        )

    doc.save(str(working_docx))
    apply_svg_pictures(
        working_docx,
        resolve_svg_requests(doc, svg_requests),
        holder_donor=holder_donor,
        caption_donor=caption_donor,
    )
    insert_formula_omml(working_docx, FORMULA_ANCHOR, "00ABC123")
    patch_formula(working_docx, "00ABC123", f"({formula_chapter_index + 1}-1)", None)
    for table_index in rendered_table_indices(working_docx):
        patch_table(working_docx, table_index, "wps_second_three_line_rendered")
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "normalize_thesis_citation_chain.py"),
            "--docx",
            str(working_docx),
        ],
        timeout=600,
    )
    bind_body_citation_paragraphs_to_style(working_docx, body_donor.style_id or "Normal")
    return working_docx, asset_manifest


def run_audits(
    template_path: Path,
    final_docx: Path,
    final_pdf: Path,
    asset_manifest: Path,
    run_root: Path,
    *,
    smoke_acceptance: bool,
) -> dict[str, Path]:
    reports_dir = run_root / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    citation_audit = reports_dir / "citation-audit.md"
    font_audit = reports_dir / "font-audit.md"
    body_style_audit = reports_dir / "body-style-audit.md"
    self_check = reports_dir / "sample-self-check.md"
    acceptance_record = reports_dir / "acceptance-record.md"
    hardfield_inputs = {
        "surface_geometry": reports_dir / "surface-geometry-required.json",
        "surface_paragraph_typography": reports_dir / "surface-paragraph-typography-required.json",
        "toc_geometry": reports_dir / "toc-geometry-required.json",
        "toc_paragraph_typography": reports_dir / "toc-paragraph-typography-required.json",
        "whole_pagination": reports_dir / "whole-pagination-required.json",
    }
    for path in hardfield_inputs.values():
        path.write_text(
            json.dumps(
                {
                    "verdict": "blocked",
                    "reason": "measured hard-field producer has not supplied this required acceptance input",
                },
                ensure_ascii=False,
                indent=2,
            )
            + "\n",
            encoding="utf-8",
        )

    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_thesis_citations.py"),
            "--docx",
            str(final_docx),
            "--report",
            str(citation_audit),
        ],
        check=False,
        timeout=600,
    )
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_docx_font_encoding.py"),
            str(final_docx),
            "--reference-docx",
            str(template_path),
            "--report",
            str(font_audit),
        ],
        check=False,
        timeout=600,
    )
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "audit_docx_body_style.py"),
            "--reference-docx",
            str(template_path),
            "--final-docx",
            str(final_docx),
            "--report",
            str(body_style_audit),
        ],
        check=False,
        timeout=600,
    )
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "sample_self_check.py"),
            "--reference-docx",
            str(template_path),
            "--final-docx",
            str(final_docx),
            "--final-pdf",
            str(final_pdf),
            "--citation-audit",
            str(citation_audit),
            "--font-audit",
            str(font_audit),
            "--body-style-audit",
            str(body_style_audit),
            "--asset-manifest",
            str(asset_manifest),
            "--output",
            str(self_check),
            "--fail-on-issues",
            *(["--smoke-acceptance"] if smoke_acceptance else []),
        ],
        timeout=1800,
    )
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "generate_thesis_acceptance_record.py"),
            "--template",
            str(template_path),
            "--source-docx",
            str(template_path),
            "--final-docx",
            str(final_docx),
            "--final-pdf",
            str(final_pdf),
            "--self-check",
            str(self_check),
            "--citation-audit",
            str(citation_audit),
            "--font-audit",
            str(font_audit),
            "--body-style-audit",
            str(body_style_audit),
            "--surface-geometry-json",
            str(hardfield_inputs["surface_geometry"]),
            "--surface-paragraph-typography-json",
            str(hardfield_inputs["surface_paragraph_typography"]),
            "--toc-geometry-json",
            str(hardfield_inputs["toc_geometry"]),
            "--toc-paragraph-typography-json",
            str(hardfield_inputs["toc_paragraph_typography"]),
            "--whole-pagination-json",
            str(hardfield_inputs["whole_pagination"]),
            "--validator",
            f"{PYTHON_EXE} {SCRIPT_DIR / 'validate_skill_gate.py'}",
            "--selftest-command",
            f"{PYTHON_EXE} {SCRIPT_DIR / 'selftest_skill_flow.py'} --include-integration",
            "--helper-scripts-planned",
            "build_minimal_template_thesis.py",
            "--delegated-canonical-helper-paths",
            str((SCRIPT_DIR / "build_minimal_template_thesis.py").resolve()),
            *(["--smoke-acceptance"] if smoke_acceptance else []),
            "--output",
            str(acceptance_record),
        ],
        timeout=1800,
    )
    if not smoke_acceptance:
        run(
            [
                str(PYTHON_EXE),
                str(SCRIPT_DIR / "validate_skill_gate.py"),
                "--gate-record",
                str(acceptance_record),
            ],
            timeout=1800,
        )
    return {
        "citation_audit": citation_audit,
        "font_audit": font_audit,
        "body_style_audit": body_style_audit,
        "self_check": self_check,
        "acceptance_record": acceptance_record,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template")
    parser.add_argument("--output-docx")
    parser.add_argument("--output-pdf")
    parser.add_argument("--run-root")
    parser.add_argument("--smoke-acceptance", action="store_true")
    parser.add_argument("--word-export-only", action="store_true", help=argparse.SUPPRESS)
    parser.add_argument("--word-export-docx", help=argparse.SUPPRESS)
    parser.add_argument("--word-export-pdf", help=argparse.SUPPRESS)
    args = parser.parse_args()

    if args.word_export_only:
        if not args.word_export_docx or not args.word_export_pdf:
            raise SystemExit("--word-export-only requires --word-export-docx and --word-export-pdf")
        export_with_word_direct(Path(args.word_export_docx).resolve(), Path(args.word_export_pdf).resolve())
        return 0

    if not args.template or not args.output_docx or not args.output_pdf:
        parser.error("--template, --output-docx, and --output-pdf are required")

    template_path = Path(args.template).resolve()
    output_docx = Path(args.output_docx).resolve()
    output_pdf = Path(args.output_pdf).resolve()
    run_root = (
        Path(args.run_root).resolve()
        if args.run_root
        else output_docx.parent / f"gpb_minimal_template_{time.strftime('%Y%m%d_%H%M%S')}"
    )
    run_root.mkdir(parents=True, exist_ok=True)
    (run_root / "meta").mkdir(parents=True, exist_ok=True)
    (run_root / "pages").mkdir(parents=True, exist_ok=True)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    working_docx = run_root / "working.docx"
    working_docx, asset_manifest = build_minimal_doc(template_path, working_docx, run_root)
    try:
        sync_static_toc_if_possible(working_docx, output_docx, run_root, reference_docx=template_path)
    except RuntimeError as exc:
        if not args.smoke_acceptance:
            raise
        shutil.copy2(working_docx, output_docx)
        (run_root / "reports").mkdir(parents=True, exist_ok=True)
        (run_root / "reports" / "smoke-toc-sync-fallback.md").write_text(
            "- smoke acceptance fallback: static TOC sync failed; copied working DOCX for non-deliverable detector fixture\n"
            f"- reason: {exc}\n",
            encoding="utf-8",
        )
    export_pdf(output_docx, output_pdf)
    run(
        [
            str(PYTHON_EXE),
            str(SCRIPT_DIR / "pdf_to_pages.py"),
            str(output_pdf),
            str(run_root / "pages"),
        ],
        check=False,
        timeout=1800,
    )
    reports = run_audits(
        template_path,
        output_docx,
        output_pdf,
        asset_manifest,
        run_root,
        smoke_acceptance=args.smoke_acceptance,
    )

    print(f"template={template_path}")
    print(f"output_docx={output_docx}")
    print(f"output_pdf={output_pdf}")
    print(f"run_root={run_root}")
    print(f"asset_manifest={asset_manifest}")
    for key, path in reports.items():
        print(f"{key}={path}")
    return 0


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    raise SystemExit(main())
