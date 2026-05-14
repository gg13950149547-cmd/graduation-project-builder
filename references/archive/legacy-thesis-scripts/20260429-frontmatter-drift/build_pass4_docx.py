from __future__ import annotations

import argparse
import gc
import json
import re
import subprocess
import shutil
import sys
import time
from pathlib import Path

from align_target_thesis import patch_section_xml
from PIL import Image
from docx import Document
from docx.document import Document as _Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


def zh(value: str) -> str:
    return value


def has_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in text)


def set_run_font(run, *, east_asia: str | None = None, size: float = 12.0, bold: bool = False):
    east = east_asia or ("SimSun" if has_cjk(run.text) else "Times New Roman")
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:cs"), "Times New Roman")
    color = rpr.color
    if color is None:
        color = OxmlElement("w:color")
        rpr.append(color)
    color.set(qn("w:val"), "000000")


def reset_paragraph(paragraph: Paragraph):
    paragraph.clear()
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = None
    pf.right_indent = None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    east_asia: str | None = None,
    size: float = 12.0,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_pt: float = 24.0,
    line_spacing: float = 1.3,
):
    reset_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, size=size, bold=bold)
    paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(first_line_indent_pt)
    pf.line_spacing = line_spacing


def set_caption(paragraph: Paragraph, text: str):
    set_paragraph_text(
        paragraph,
        text,
        east_asia="SimSun",
        size=12.0,
        bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0.0,
        line_spacing=1.0,
    )
    paragraph.paragraph_format.space_after = Pt(6)


def delete_paragraph(paragraph: Paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_paragraph_before_table(table: Table) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._element.addprevious(new_p)
    return Paragraph(new_p, table._parent)


def paragraph_position(doc: Document, target: Paragraph) -> int:
    for idx, para in enumerate(doc.paragraphs, start=1):
        if para._element is target._element:
            return idx
    raise KeyError("Paragraph not found in document")


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def is_heading_text(text: str) -> bool:
    return bool(re.match(r"^\d+(\.\d+){0,2}\s", text)) or text in {
        zh("\u7ed3   \u8bba"),
        zh("\u81f4   \u8c22"),
        zh("\u53c2\u8003\u6587\u732e"),
    }


def find_paragraph_by_prefix(doc: Document, prefix: str) -> Paragraph:
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") and para.text.strip().startswith(prefix):
            return para
    raise KeyError(prefix)


def last_body_paragraph_in_section(doc: Document, heading_prefix: str) -> Paragraph:
    heading_index = None
    heading = None
    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") and para.text.strip().startswith(heading_prefix):
            heading_index = idx
            heading = para
            break
    if heading is None or heading_index is None:
        raise KeyError(heading_prefix)
    last = heading
    for para in doc.paragraphs[heading_index + 1 :]:
        text = para.text.strip()
        if text and is_heading_text(text):
            break
        if text:
            last = para
    return last


def add_picture_block(
    anchor: Paragraph,
    intro: str,
    image_path: Path,
    caption: str,
    summary: str,
    width_cm: float,
    *,
    svg_path: Path | None = None,
    svg_fallback: Path | None = None,
    pending_svg: list[dict] | None = None,
) -> Paragraph:
    intro_para = insert_paragraph_after(anchor)
    set_paragraph_text(intro_para, intro)

    img_para = insert_paragraph_after(intro_para)
    reset_paragraph(img_para)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.keep_with_next = True
    if svg_path is None or pending_svg is None:
        img_para.add_run().add_picture(str(image_path), width=Cm(width_cm))
    else:
        pending_svg.append(
            {
                "paragraph": img_para,
                "svg_path": svg_path,
                "fallback_png": svg_fallback or image_path,
                "width_cm": width_cm,
            }
        )

    cap_para = insert_paragraph_after(img_para)
    set_caption(cap_para, caption)

    sum_para = insert_paragraph_after(cap_para)
    set_paragraph_text(sum_para, summary)
    return sum_para


def format_cover_line(paragraph: Paragraph, label: str, value: str):
    reset_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)

    label_run = paragraph.add_run(label)
    set_run_font(label_run, east_asia="SimSun", size=14.0, bold=False)

    value_run = paragraph.add_run(f"{value}{' ' * (20 if value else 26)}")
    set_run_font(value_run, east_asia="SimSun", size=14.0, bold=False)
    value_run.underline = True


def rebuild_cover(doc: Document, assets: dict):
    anchor = doc.paragraphs[10]
    for para in list(doc.paragraphs[:10]):
        delete_paragraph(para)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    p = anchor.insert_paragraph_before()
    reset_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.add_run().add_picture(assets["reference_media"]["image1.jpeg"], width=Cm(6.4))

    p = anchor.insert_paragraph_before()
    set_paragraph_text(
        p,
        zh("\u6bd5 \u4e1a \u8bbe \u8ba1"),
        east_asia="SimHei",
        size=28.0,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0.0,
        line_spacing=1.0,
    )
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(18)

    p = anchor.insert_paragraph_before()
    set_paragraph_text(
        p,
        zh("\u9898\u76ee\uff1a\u5b66\u751f\u8bfe\u5802\u884c\u4e3a\u68c0\u6d4b\u4e0e\u8bc4\u4f30\u7cfb\u7edf\u7684"),
        east_asia="SimSun",
        size=19.0,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0.0,
        line_spacing=1.0,
    )
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

    p = anchor.insert_paragraph_before()
    set_paragraph_text(
        p,
        zh("\u8bbe\u8ba1\u4e0e\u5b9e\u73b0"),
        east_asia="SimSun",
        size=19.0,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0.0,
        line_spacing=1.0,
    )
    p.paragraph_format.space_after = Pt(22)

    p = anchor.insert_paragraph_before()
    reset_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(assets["reference_media"]["image2.jpeg"], width=Cm(5.4))
    p.paragraph_format.space_after = Pt(18)

    cover_rows = [
        (zh("\u59d3        \u540d\uff1a"), zh("\u8d75\u632f\u7ff0")),
        (zh("\u5b66        \u53f7\uff1a"), "2022220360005"),
        (zh("\u5b66        \u9662\uff1a"), zh("\u673a\u5668\u4eba\u5b66\u9662")),
        (zh("\u4e13        \u4e1a\uff1a"), zh("\u8f6f\u4ef6\u5de5\u7a0b")),
        (zh("\u6307 \u5bfc  \u6559 \u5e08\uff1a"), zh("\u6881\u6654")),
        (zh("\u534f\u52a9\u6307\u5bfc\u6559\u5e08\uff1a"), ""),
    ]
    for label, value in cover_rows:
        p = anchor.insert_paragraph_before()
        format_cover_line(p, label, value)

    p = anchor.insert_paragraph_before()
    set_paragraph_text(
        p,
        zh("2026 \u5e74 4 \u6708 18 \u65e5"),
        east_asia="SimSun",
        size=16.0,
        bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0.0,
        line_spacing=1.0,
    )
    p.paragraph_format.space_before = Pt(18)


def resolve_svg_requests(doc: Document, pending_svg: list[dict]) -> list[dict]:
    resolved: list[dict] = []
    for item in pending_svg:
        resolved.append(
            {
                "paragraph_index": paragraph_position(doc, item["paragraph"]),
                "svg_path": Path(item["svg_path"]),
                "fallback_png": Path(item["fallback_png"]),
                "width_cm": float(item["width_cm"]),
            }
        )
    return resolved


def svg_height_cm(fallback_png: Path, width_cm: float) -> float:
    with Image.open(fallback_png) as img:
        return width_cm * img.height / img.width


def apply_svg_pictures(docx_path: Path, requests: list[dict]) -> None:
    for item in requests:
        height_cm = svg_height_cm(item["fallback_png"], item["width_cm"])
        command = [
            "officecli",
            "add",
            str(docx_path),
            f"/body/p[{item['paragraph_index']}]",
            "--type",
            "picture",
            "--prop",
            f"src={item['svg_path']}",
            "--prop",
            f"fallback={item['fallback_png']}",
            "--prop",
            f"width={item['width_cm']:.2f}cm",
            "--prop",
            f"height={height_cm:.2f}cm",
        ]
        last_error: subprocess.CalledProcessError | subprocess.TimeoutExpired | None = None
        for attempt in range(4):
            try:
                subprocess.run(command, check=True, timeout=120)
                last_error = None
                break
            except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
                last_error = exc
                time.sleep(0.4 * (attempt + 1))
        if last_error is not None:
            raise last_error


def insert_figure_family(doc: Document, assets: dict) -> list[dict]:
    d = assets["diagrams"]
    s = assets["screenshots"]
    pending_svg: list[dict] = []

    anchor = last_body_paragraph_in_section(doc, "4.1 ")
    anchor = add_picture_block(
        anchor,
        zh("\u5982\u56fe 1 \u6240\u793a\uff0c\u7cfb\u7edf\u603b\u4f53\u67b6\u6784\u56f4\u7ed5\u524d\u7aef\u4ea4\u4e92\uff0cFlask \u540e\u7aef\u63a5\u53e3\uff0cFFmpeg \u89c6\u9891\u5904\u7406\uff0cYOLOv8n-pose \u884c\u4e3a\u8bc6\u522b\u4e0e MySQL/Redis \u5b58\u50a8\u94fe\u8def\u5c55\u5f00\u3002\u56fe\u4e2d\u7528\u4ece\u4e0a\u5230\u4e0b\u7684\u5c42\u6b21\u8868\u793a\u8bf7\u6c42\u5904\u7406\uff0c\u7528\u4ece\u5de6\u5230\u53f3\u7684\u8fde\u7ebf\u8868\u793a\u6570\u636e\u6d41\u5411\uff0c\u5bf9\u5e94\u4e86\u7cfb\u7edf\u5b9e\u9645\u7684\u524d\u540e\u7aef\u5206\u5c42\u5b9e\u73b0\u3002"),
        Path(d["fig01_system_architecture"]["png"]),
        zh("\u56fe 1 \u7cfb\u7edf\u603b\u4f53\u67b6\u6784\u56fe"),
        zh("\u8be5\u67b6\u6784\u4f53\u73b0\u4e86\u9875\u9762\u4ea4\u4e92\uff0c\u4e1a\u52a1\u5904\u7406\uff0c\u6a21\u578b\u63a8\u7406\u548c\u6570\u636e\u5b58\u50a8\u7684\u804c\u8d23\u5206\u79bb\uff0c\u4e5f\u4e3a\u540e\u7eed\u5bf9\u6a21\u5757\u8bbe\u8ba1\uff0c\u8fd0\u884c\u6d41\u7a0b\u548c\u6570\u636e\u7ed3\u6784\u7684\u8bba\u8ff0\u63d0\u4f9b\u4e86\u7edf\u4e00\u7684\u7ec4\u7ec7\u57fa\u7840\u3002"),
        15.2,
        svg_path=Path(d["fig01_system_architecture"]["svg"]),
        svg_fallback=Path(d["fig01_system_architecture"]["png"]),
        pending_svg=pending_svg,
    )

    anchor = last_body_paragraph_in_section(doc, "4.2 ")
    anchor = add_picture_block(
        anchor,
        zh("\u5982\u56fe 2 \u6240\u793a\uff0c\u7cfb\u7edf\u603b\u4f53\u529f\u80fd\u6a21\u5757\u4ee5\u89c6\u9891\u4e0a\u4f20\uff0c\u884c\u4e3a\u68c0\u6d4b\uff0c\u8bfe\u5802\u8bc4\u4f30\uff0c\u6570\u636e\u7edf\u8ba1\u548c\u5386\u53f2\u56de\u653e\u4e3a\u4e3b\u7ebf\uff0c\u540c\u65f6\u7531\u7528\u6237\u6743\u9650\u7ba1\u7406\u6a21\u5757\u8d1f\u8d23\u4e0d\u540c\u89d2\u8272\u7684\u8bbf\u95ee\u63a7\u5236\u3002\u5404\u6a21\u5757\u7684\u5206\u5de5\u6e05\u6670\uff0c\u80fd\u591f\u76f4\u89c2\u652f\u6491\u540e\u6587\u7684\u529f\u80fd\u5b9e\u73b0\u7ae0\u8282\u5c55\u5f00\u3002"),
        Path(d["fig02_function_modules"]["png"]),
        zh("\u56fe 2 \u7cfb\u7edf\u603b\u4f53\u529f\u80fd\u6a21\u5757\u56fe"),
        zh("\u91c7\u7528\u6a21\u5757\u5316\u7ed3\u6784\u540e\uff0c\u7cfb\u7edf\u53ef\u4ee5\u5728\u4e0d\u7834\u574f\u4e3b\u6d41\u7a0b\u7684\u60c5\u51b5\u4e0b\u5bf9\u5355\u4e2a\u529f\u80fd\u8fdb\u884c\u6269\u5c55\u6216\u4f18\u5316\uff0c\u4e5f\u4f7f\u8bba\u6587\u7684\u7b2c 5 \u7ae0\u80fd\u591f\u6309\u7167\u529f\u80fd\u8fb9\u754c\u8fdb\u884c\u7ec6\u5316\u63cf\u8ff0\u3002"),
        15.0,
        svg_path=Path(d["fig02_function_modules"]["svg"]),
        svg_fallback=Path(d["fig02_function_modules"]["png"]),
        pending_svg=pending_svg,
    )

    anchor = last_body_paragraph_in_section(doc, "4.3.1 ")
    anchor = add_picture_block(
        anchor,
        zh("\u5982\u56fe 3 \u6240\u793a\uff0c\u6570\u636e\u5e93 E-R \u7ed3\u6784\u4ee5\u7528\u6237\uff0c\u89c6\u9891\uff0c\u68c0\u6d4b\u4efb\u52a1\uff0c\u68c0\u6d4b\u7ed3\u679c\u548c\u8bc4\u4f30\u62a5\u544a\u4e94\u7c7b\u5b9e\u4f53\u4e3a\u6838\u5fc3\uff0c\u8986\u76d6\u4e86\u7cfb\u7edf\u4ece\u89c6\u9891\u4e0a\u4f20\uff0c\u4efb\u52a1\u6267\u884c\uff0c\u7ed3\u679c\u5165\u5e93\u5230\u8bc4\u4f30\u62a5\u544a\u751f\u6210\u7684\u5b8c\u6574\u6570\u636e\u94fe\u8def\u3002"),
        Path(d["fig03_er_model"]["png"]),
        zh("\u56fe 3 \u6570\u636e\u5e93 E-R \u56fe"),
        zh("\u8fd9\u4e00\u7ed3\u6784\u4e0d\u4ec5\u7528\u4e8e\u652f\u6491\u5f53\u524d\u8bba\u6587\u4e2d\u7684\u529f\u80fd\u5b9e\u73b0\uff0c\u4e5f\u4e3a\u540e\u7eed\u7684\u5386\u53f2\u5206\u6790\uff0c\u8bc4\u4f30\u62a5\u544a\u67e5\u8be2\u548c\u6570\u636e\u7edf\u8ba1\u63d0\u4f9b\u4e86\u53ef\u8ffd\u6eaf\u7684\u4e3b\u952e\u4e0e\u5916\u952e\u5173\u7cfb\u57fa\u7840\u3002"),
        15.0,
        svg_path=Path(d["fig03_er_model"]["svg"]),
        svg_fallback=Path(d["fig03_er_model"]["png"]),
        pending_svg=pending_svg,
    )

    anchor = last_body_paragraph_in_section(doc, "5.2.1 ")
    anchor = add_picture_block(
        anchor,
        zh("\u56fe 4 \u5bf9\u5e94\u89c6\u9891\u4e0a\u4f20\u7684\u672c\u5730\u6587\u4ef6\u9009\u62e9\u73af\u8282\u3002\u7528\u6237\u5728\u7cfb\u7edf\u89c6\u9891\u7ba1\u7406\u9875\u8c03\u7528\u6587\u4ef6\u9009\u62e9\u5bf9\u8bdd\u6846\u540e\uff0c\u53ef\u4ee5\u6309\u7167\u8bfe\u5802\u6444\u50cf\u8d44\u6e90\u7684\u5b58\u50a8\u76ee\u5f55\u5b9a\u4f4d\u89c6\u9891\u3002\u8fd9\u4e2a\u754c\u9762\u76f4\u63a5\u4f53\u73b0\u4e86\u7cfb\u7edf\u652f\u6301\u771f\u5b9e\u4e1a\u52a1\u89c6\u9891\u7d20\u6750\u63a5\u5165\uff0c\u800c\u4e0d\u662f\u4ec5\u505c\u7559\u5728\u6a21\u62df\u6570\u636e\u7684\u6f14\u793a\u5c42\u9762\u3002"),
        Path(s["fig04_file_dialog"]),
        zh("\u56fe 4 \u89c6\u9891\u6587\u4ef6\u9009\u62e9\u5f39\u7a97"),
        zh("\u901a\u8fc7\u8be5\u6b65\u9aa4\uff0c\u89c6\u9891\u4e0a\u4f20\u6d41\u7a0b\u5b8c\u6210\u4e86\u4ece\u7528\u6237\u64cd\u4f5c\u5230\u4e1a\u52a1\u8d44\u6e90\u5b9a\u4f4d\u7684\u8854\u63a5\uff0c\u4e3a\u540e\u7eed\u7684\u683c\u5f0f\u6821\u9a8c\uff0c\u5143\u6570\u636e\u63d0\u53d6\u548c\u9884\u5904\u7406\u4efb\u52a1\u751f\u6210\u5960\u5b9a\u8f93\u5165\u57fa\u7840\u3002"),
        14.4,
    )
    anchor = add_picture_block(
        anchor,
        zh("\u56fe 5 \u5bf9\u5e94\u89c6\u9891\u4e0a\u4f20\u5b8c\u6210\u540e\u7684\u7ba1\u7406\u754c\u9762\u3002\u9875\u9762\u4e0d\u4ec5\u663e\u793a\u5df2\u4e0a\u4f20\u7684\u89c6\u9891\u6587\u4ef6\u540d\uff0c\u8fd8\u540c\u65f6\u5c55\u793a\u65f6\u957f\uff0c\u5206\u8fa8\u7387\uff0cFPS\uff0c\u603b\u5e27\u6570\uff0c\u4e0a\u4f20\u8005\uff0c\u72b6\u6001\u548c\u4e0a\u4f20\u65f6\u95f4\u7b49\u4fe1\u606f\uff0c\u4f7f\u7528\u6237\u53ef\u4ee5\u5728\u540c\u4e00\u4e2a\u9875\u9762\u4e2d\u5b8c\u6210\u4efb\u52a1\u72b6\u6001\u786e\u8ba4\u3002"),
        Path(s["fig05_video_management"]),
        zh("\u56fe 5 \u89c6\u9891\u4e0a\u4f20\u4e0e\u7ba1\u7406\u754c\u9762"),
        zh("\u8fd9\u79cd\u5c06\u4e0a\u4f20\u5165\u53e3\uff0c\u6587\u4ef6\u4fe1\u606f\u548c\u5904\u7406\u72b6\u6001\u96c6\u4e2d\u5448\u73b0\u7684\u65b9\u5f0f\uff0c\u80fd\u591f\u6709\u6548\u964d\u4f4e\u7528\u6237\u5728\u4e0d\u540c\u9875\u9762\u4e4b\u95f4\u6765\u56de\u786e\u8ba4\u4efb\u52a1\u8fdb\u5ea6\u7684\u64cd\u4f5c\u6210\u672c\u3002"),
        15.2,
    )

    anchor = last_body_paragraph_in_section(doc, "5.2.2 ")
    anchor = add_picture_block(
        anchor,
        zh("\u5982\u56fe 6 \u6240\u793a\uff0c\u89c6\u9891\u5904\u7406\u6d41\u7a0b\u4f9d\u6b21\u5305\u62ec\u4e0a\u4f20\uff0c\u683c\u5f0f\u6821\u9a8c\uff0c\u5143\u6570\u636e\u63d0\u53d6\uff0c\u8f6c\u7801\u4e0e\u62bd\u5e27\uff0c\u4efb\u52a1\u751f\u6210\u53ca\u72b6\u6001\u5199\u56de\u3002\u8be5\u6d41\u7a0b\u628a\u539f\u59cb\u89c6\u9891\u8f6c\u6362\u6210\u540e\u7eed\u68c0\u6d4b\u6a21\u5757\u53ef\u76f4\u63a5\u4f7f\u7528\u7684\u8f93\u5165\u6570\u636e\u5bf9\u8c61\uff0c\u540c\u65f6\u628a\u5904\u7406\u8fdb\u5ea6\u4fdd\u5b58\u5230\u540e\u7aef\u8bb0\u5f55\u4e2d\u3002"),
        Path(d["fig06_video_processing_flow"]["png"]),
        zh("\u56fe 6 \u89c6\u9891\u5904\u7406\u6d41\u7a0b\u56fe"),
        zh("\u89c6\u9891\u9884\u5904\u7406\u6a21\u5757\u7684\u4ef7\u503c\u5728\u4e8e\u5b83\u628a\u7528\u6237\u4e0a\u4f20\u7684\u539f\u59cb\u8d44\u6599\u8f6c\u5316\u6210\u68c0\u6d4b\u5f15\u64ce\u53ef\u8bc6\u522b\u7684\u6807\u51c6\u5316\u6570\u636e\uff0c\u4ece\u800c\u4fdd\u8bc1\u540e\u7eed\u6d41\u6c34\u7ebf\u7684\u8f93\u5165\u4e00\u81f4\u6027\u548c\u7ed3\u679c\u53ef\u6bd4\u6027\u3002"),
        15.0,
        svg_path=Path(d["fig06_video_processing_flow"]["svg"]),
        svg_fallback=Path(d["fig06_video_processing_flow"]["png"]),
        pending_svg=pending_svg,
    )

    anchor = last_body_paragraph_in_section(doc, "5.3.3 ")
    anchor = add_picture_block(
        anchor,
        zh("\u5982\u56fe 7 \u6240\u793a\uff0c\u884c\u4e3a\u68c0\u6d4b\u6d41\u7a0b\u4ece\u89c6\u9891\u62bd\u5e27\u5f00\u59cb\uff0c\u518d\u901a\u8fc7 YOLOv8n-pose \u8fdb\u884c\u5173\u952e\u70b9\u68c0\u6d4b\uff0c\u5728\u6b64\u57fa\u7840\u4e0a\u6839\u636e\u89c4\u5219\u5b8c\u6210\u542c\u8bb2\uff0c\u4e3e\u624b\uff0c\u8f6c\u5934\uff0c\u4f4e\u5934\uff0c\u8db4\u684c\u548c\u8bfb\u5199\u7b49\u884c\u4e3a\u5224\u5b9a\uff0c\u6700\u540e\u628a\u7ed3\u679c\u6279\u91cf\u5199\u5165\u6570\u636e\u5e93\u4f9b\u524d\u7aef\u67e5\u8be2\u3002"),
        Path(d["fig07_detection_pipeline_flow"]["png"]),
        zh("\u56fe 7 \u884c\u4e3a\u68c0\u6d4b\u6d41\u7a0b\u56fe"),
        zh("\u8be5\u6d41\u7a0b\u56fe\u628a\u7b97\u6cd5\u903b\u8f91\u4e0e\u5de5\u7a0b\u903b\u8f91\u7edf\u4e00\u8d77\u6765\uff0c\u8bf4\u660e\u7cfb\u7edf\u4e0d\u4ec5\u80fd\u505a\u79bb\u7ebf\u6a21\u578b\u63a8\u7406\uff0c\u8fd8\u80fd\u5728\u4efb\u52a1\u7ea7\u522b\u4fdd\u6301\u7a33\u5b9a\u7684\u5165\u5e93\uff0c\u65e5\u5fd7\u548c\u53ef\u89c6\u5316\u53cd\u9988\u80fd\u529b\u3002"),
        15.0,
        svg_path=Path(d["fig07_detection_pipeline_flow"]["svg"]),
        svg_fallback=Path(d["fig07_detection_pipeline_flow"]["png"]),
        pending_svg=pending_svg,
    )
    anchor = add_picture_block(
        anchor,
        zh("\u56fe 8 \u5c55\u793a\u4e86\u68c0\u6d4b\u6d41\u6c34\u7ebf\u5728\u5b9e\u9645\u9875\u9762\u4e2d\u7684\u5448\u73b0\u65b9\u5f0f\u3002\u5de6\u4fa7\u4efb\u52a1\u5217\u8868\u8bf4\u660e\u7528\u6237\u53ef\u4ee5\u4ece\u5df2\u4e0a\u4f20\u89c6\u9891\u4e2d\u542f\u52a8\u68c0\u6d4b\u4efb\u52a1\uff0c\u53f3\u4fa7\u5206\u6790\u7a97\u53e3\u5219\u7ed9\u51fa\u4e86\u5404\u7c7b\u884c\u4e3a\u7684\u5206\u5e03\u5360\u6bd4\u548c\u65f6\u5e8f\u53d8\u5316\uff0c\u8bf4\u660e\u7cfb\u7edf\u80fd\u628a\u540e\u53f0\u68c0\u6d4b\u7ed3\u679c\u53ca\u65f6\u8f6c\u5316\u4e3a\u7528\u6237\u53ef\u7406\u89e3\u7684\u5206\u6790\u754c\u9762\u3002"),
        Path(s["fig08_detection_task_analysis"]),
        zh("\u56fe 8 \u68c0\u6d4b\u4efb\u52a1\u4e0e\u7ed3\u679c\u5206\u6790\u754c\u9762"),
        zh("\u4ece\u5b9e\u73b0\u6548\u679c\u6765\u770b\uff0c\u68c0\u6d4b\u6d41\u6c34\u7ebf\u5df2\u7ecf\u5f62\u6210\u4ece\u4efb\u52a1\u521b\u5efa\uff0c\u8fdb\u5ea6\u53ef\u89c6\uff0c\u7ed3\u679c\u67e5\u770b\u5230\u6570\u636e\u7edf\u8ba1\u7684\u95ed\u73af\uff0c\u8fd9\u4e5f\u662f\u8bba\u6587\u4e2d\u884c\u4e3a\u68c0\u6d4b\u6a21\u5757\u5de5\u7a0b\u5b9e\u7528\u6027\u7684\u76f4\u89c2\u4f53\u73b0\u3002"),
        15.2,
    )

    anchor = last_body_paragraph_in_section(doc, "5.4.2 ")
    anchor = add_picture_block(
        anchor,
        zh("\u5982\u56fe 9 \u6240\u793a\uff0c\u8bfe\u5802\u8d28\u91cf\u8bc4\u4f30\u6d41\u7a0b\u4ee5\u68c0\u6d4b\u7ed3\u679c\u6570\u636e\u4e3a\u8f93\u5165\uff0c\u901a\u8fc7\u65f6\u95f4\u7a97\u53e3\u805a\u5408\uff0c\u6307\u6807\u8ba1\u7b97\uff0c\u4e13\u6ce8\u5ea6\u8bc4\u5206\uff0c\u7b49\u7ea7\u5224\u5b9a\u548c\u62a5\u544a\u751f\u6210\u7b49\u73af\u8282\uff0c\u5c06\u539f\u59cb\u884c\u4e3a\u6570\u636e\u8f6c\u5316\u6210\u4fbf\u4e8e\u6559\u5b66\u89e3\u8bfb\u7684\u8bc4\u4f30\u7ed3\u8bba\u3002"),
        Path(d["fig09_evaluation_flow"]["png"]),
        zh("\u56fe 9 \u8bfe\u5802\u8d28\u91cf\u8bc4\u4f30\u6d41\u7a0b\u56fe"),
        zh("\u8bc4\u4f30\u6a21\u5757\u7684\u4f5c\u7528\u5728\u4e8e\u5c06\u79bb\u6563\u7684\u68c0\u6d4b\u7ed3\u679c\u91cd\u65b0\u7ec4\u7ec7\u4e3a\u6307\u6807\uff0c\u5206\u7ea7\uff0c\u56fe\u8868\u548c\u6587\u5b57\u5efa\u8bae\uff0c\u4f7f\u7cfb\u7edf\u80fd\u591f\u4ece\u201c\u68c0\u6d4b\u201d\u7ee7\u7eed\u8d70\u5411\u201c\u8bc4\u4f30\u201d\u3002"),
        15.0,
        svg_path=Path(d["fig09_evaluation_flow"]["svg"]),
        svg_fallback=Path(d["fig09_evaluation_flow"]["png"]),
        pending_svg=pending_svg,
    )
    anchor = add_picture_block(
        anchor,
        zh("\u56fe 10 \u4e3a\u8bc4\u4f30\u62a5\u544a\u8be6\u60c5\u754c\u9762\u3002\u754c\u9762\u5728\u5c55\u793a\u4e13\u6ce8\u5ea6\u5f97\u5206\u7684\u540c\u65f6\uff0c\u8fd8\u7ed9\u51fa\u5404\u7c7b\u884c\u4e3a\u5360\u6bd4\u56fe\u548c\u6559\u5b66\u5efa\u8bae\uff0c\u4f7f\u6559\u5e08\u6216\u7ba1\u7406\u8005\u4e0d\u9700\u8981\u91cd\u65b0\u56de\u770b\u5e95\u5c42\u68c0\u6d4b\u8bb0\u5f55\uff0c\u5373\u53ef\u76f4\u63a5\u5bf9\u8be5\u8282\u8bfe\u7684\u8bfe\u5802\u72b6\u6001\u4f5c\u51fa\u5224\u65ad\u3002"),
        Path(s["fig10_evaluation_report"]),
        zh("\u56fe 10 \u8bfe\u5802\u8bc4\u4f30\u62a5\u544a\u8be6\u60c5\u754c\u9762"),
        zh("\u8fd9\u4e00\u754c\u9762\u8bf4\u660e\u7cfb\u7edf\u5df2\u5177\u5907\u4ece\u6307\u6807\u8ba1\u7b97\u5230\u7ed3\u679c\u5c55\u793a\u7684\u5b8c\u6574\u8f93\u51fa\u80fd\u529b\uff0c\u4e5f\u8ba9\u8bba\u6587\u7b2c 5.4 \u8282\u7684\u201c\u62a5\u544a\u751f\u6210\u201d\u90e8\u5206\u6709\u4e86\u660e\u786e\u7684\u5b9e\u9645\u652f\u6491\u754c\u9762\u3002"),
        13.2,
    )

    anchor = last_body_paragraph_in_section(doc, "5.5 ")
    anchor = add_picture_block(
        anchor,
        zh("\u56fe 11 \u5c55\u793a\u4e86\u6570\u636e\u7edf\u8ba1\u4e0e\u53ef\u89c6\u5316\u754c\u9762\u7684\u5b9e\u9645\u6548\u679c\u3002\u9875\u9762\u901a\u8fc7\u4e13\u6ce8\u5ea6\u8d8b\u52bf\u56fe\uff0c\u8bc4\u5206\u5206\u5e03\u56fe\uff0c\u884c\u4e3a\u5747\u503c\u56fe\u548c\u8fd1\u671f\u8bfe\u5802\u5bf9\u6bd4\u56fe\u5bf9\u8bfe\u5802\u6570\u636e\u8fdb\u884c\u7ec4\u5408\u5c55\u793a\uff0c\u4f7f\u7528\u6237\u80fd\u591f\u5728\u4e00\u4e2a\u754c\u9762\u4e2d\u540c\u65f6\u89c2\u5bdf\u6574\u4f53\u6c34\u5e73\uff0c\u884c\u4e3a\u6784\u6210\u548c\u4e0d\u540c\u8bfe\u6b21\u4e4b\u95f4\u7684\u5dee\u5f02\u3002"),
        Path(s["fig11_statistics_dashboard"]),
        zh("\u56fe 11 \u6570\u636e\u7edf\u8ba1\u4e0e\u53ef\u89c6\u5316\u754c\u9762"),
        zh("\u76f8\u6bd4\u5355\u4e00\u7684\u7edf\u8ba1\u8868\u683c\uff0c\u8be5\u754c\u9762\u66f4\u5f3a\u8c03\u6570\u636e\u7684\u8d8b\u52bf\u6027\u548c\u6bd4\u8f83\u6027\uff0c\u80fd\u4e3a\u8bfe\u5802\u8d28\u91cf\u8ddf\u8e2a\uff0c\u5f02\u5e38\u8bfe\u5802\u8bc6\u522b\u548c\u7ba1\u7406\u51b3\u7b56\u63d0\u4f9b\u66f4\u76f4\u63a5\u7684\u6570\u636e\u4f9d\u636e\u3002"),
        15.2,
    )

    anchor = last_body_paragraph_in_section(doc, "5.6 ")
    add_picture_block(
        anchor,
        zh("\u56fe 12 \u5c55\u793a\u4e86\u5386\u53f2\u8bb0\u5f55\u56de\u653e\u6a21\u5757\u7684\u5b9e\u9645\u4ea4\u4e92\u754c\u9762\u3002\u5de6\u4fa7\u662f\u6839\u636e\u65f6\u95f4\u5217\u51fa\u7684\u5386\u53f2\u68c0\u6d4b\u8bb0\u5f55\uff0c\u53f3\u4fa7\u5219\u540c\u6b65\u663e\u793a\u89c6\u9891\u64ad\u653e\uff0c\u5f53\u524d\u5e27\u884c\u4e3a\u6807\u7b7e\u548c\u5b9e\u65f6\u7edf\u8ba1\u4fe1\u606f\uff0c\u4f7f\u7528\u6237\u80fd\u591f\u5728\u56de\u653e\u8fc7\u7a0b\u4e2d\u5bf9\u8bfe\u5802\u72b6\u6001\u8fdb\u884c\u7ec6\u7c92\u5ea6\u590d\u76d8\u3002"),
        Path(s["fig12_history_playback"]),
        zh("\u56fe 12 \u5386\u53f2\u8bb0\u5f55\u56de\u653e\u754c\u9762"),
        zh("\u8fd9\u4e00\u6a21\u5757\u628a\u89c6\u9891\u56de\u653e\uff0c\u5386\u53f2\u8bb0\u5f55\u68c0\u7d22\u548c\u884c\u4e3a\u5206\u6790\u8fdb\u884c\u4e86\u6574\u5408\uff0c\u8bf4\u660e\u7cfb\u7edf\u4e0d\u4ec5\u80fd\u5904\u7406\u8bfe\u5802\u89c6\u9891\uff0c\u8fd8\u80fd\u4e3a\u8bfe\u540e\u7edf\u8ba1\u5206\u6790\u548c\u6559\u5b66\u590d\u76d8\u63d0\u4f9b\u53ef\u6301\u7eed\u7684\u67e5\u8be2\u652f\u6491\u3002"),
        15.2,
    )
    return pending_svg


def set_border(node, val: str, sz: str = "8"):
    node.set(qn("w:val"), val)
    node.set(qn("w:sz"), sz)
    node.set(qn("w:color"), "000000")
    node.set(qn("w:space"), "0")


def clear_descendant_shading(node) -> None:
    for shd in list(node.iter(qn("w:shd"))):
        parent = shd.getparent()
        if parent is not None:
            parent.remove(shd)


def style_table(table: Table, idx: int):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    width_map = {
        **{i: [Cm(3.0), Cm(3.0), Cm(2.5), Cm(6.0)] for i in range(1, 6)},
        **{i: [Cm(1.5), Cm(4.0), Cm(4.5), Cm(4.5)] for i in range(6, 12)},
        12: [Cm(5.0), Cm(9.5)],
    }

    tblPr = table._element.tblPr
    clear_descendant_shading(tblPr)
    tblStyle = tblPr.first_child_found_in("w:tblStyle")
    if tblStyle is not None:
        tblPr.remove(tblStyle)
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for name, val, sz in (
        ("top", "single", "8"),
        ("bottom", "single", "8"),
        ("left", "nil", "0"),
        ("right", "nil", "0"),
        ("insideH", "nil", "0"),
        ("insideV", "nil", "0"),
    ):
        node = tblBorders.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tblBorders.append(node)
        set_border(node, val, sz)

    for c_idx, width in enumerate(width_map[idx]):
        for row in table.rows:
            row.cells[c_idx].width = width

    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tcPr = cell._tc.get_or_add_tcPr()
            clear_descendant_shading(cell._tc)
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)
            for border_name in ("top", "bottom", "left", "right"):
                node = tcBorders.find(qn(f"w:{border_name}"))
                if node is None:
                    node = OxmlElement(f"w:{border_name}")
                    tcBorders.append(node)
                set_border(node, "nil", "0")
            if r_idx == 0:
                bottom = tcBorders.find(qn("w:bottom"))
                set_border(bottom, "single", "4")

            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.right_indent = Pt(0)
                if r_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif idx <= 5:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in (0, 3) else WD_ALIGN_PARAGRAPH.CENTER
                elif idx <= 11:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=12.0, bold=r_idx == 0)


def patch_sections_with_retry(docx_path: Path) -> bool:
    last_error: Exception | None = None
    for attempt in range(5):
        try:
            patch_section_xml(docx_path)
            return True
        except PermissionError as exc:
            last_error = exc
            time.sleep(0.5 * (attempt + 1))
    if last_error is not None:
        print(f"warning: patch_section_xml skipped after retry: {last_error}", file=sys.stderr)
    return False


def add_table_captions(doc: Document):
    caption_texts = {
        1: zh("\u8868 1 \u7528\u6237\u8868\u7ed3\u6784"),
        2: zh("\u8868 2 \u89c6\u9891\u8868\u7ed3\u6784"),
        3: zh("\u8868 3 \u68c0\u6d4b\u4efb\u52a1\u8868\u7ed3\u6784"),
        4: zh("\u8868 4 \u68c0\u6d4b\u7ed3\u679c\u8868\u7ed3\u6784"),
        5: zh("\u8868 5 \u8bc4\u4f30\u62a5\u544a\u8868\u7ed3\u6784"),
        6: zh("\u8868 6 \u767b\u5f55\u529f\u80fd\u6d4b\u8bd5\u7528\u4f8b\u8868"),
        7: zh("\u8868 7 \u6ce8\u518c\u529f\u80fd\u6d4b\u8bd5\u7528\u4f8b\u8868"),
        8: zh("\u8868 8 \u89c6\u9891\u4e0a\u4f20\u529f\u80fd\u6d4b\u8bd5\u7528\u4f8b\u8868"),
        9: zh("\u8868 9 \u884c\u4e3a\u68c0\u6d4b\u529f\u80fd\u6d4b\u8bd5\u7528\u4f8b\u8868"),
        10: zh("\u8868 10 \u8bfe\u5802\u8bc4\u4f30\u529f\u80fd\u6d4b\u8bd5\u7528\u4f8b\u8868"),
        11: zh("\u8868 11 \u6570\u636e\u7edf\u8ba1\u529f\u80fd\u6d4b\u8bd5\u7528\u4f8b\u8868"),
        12: zh("\u8868 12 \u7cfb\u7edf\u6027\u80fd\u6d4b\u8bd5\u7ed3\u679c\u8868"),
    }
    tables = [item for item in iter_block_items(doc) if isinstance(item, Table)]
    for idx, table in enumerate(tables, start=1):
        cap = insert_paragraph_before_table(table)
        set_caption(cap, caption_texts[idx])
        cap.paragraph_format.keep_with_next = True
        style_table(table, idx)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--base-docx", required=True)
    parser.add_argument("--asset-manifest", required=True)
    parser.add_argument("--output-docx", required=True)
    args = parser.parse_args()

    base_docx = Path(args.base_docx).resolve()
    output_docx = Path(args.output_docx).resolve()
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(base_docx, output_docx)

    assets = json.loads(Path(args.asset_manifest).read_text(encoding="utf-8"))
    doc = Document(output_docx)

    rebuild_cover(doc, assets)
    pending_svg = insert_figure_family(doc, assets)
    add_table_captions(doc)
    resolved_svg = resolve_svg_requests(doc, pending_svg)

    doc.save(output_docx)
    del doc
    gc.collect()
    time.sleep(0.2)
    apply_svg_pictures(output_docx, resolved_svg)
    patch_sections_with_retry(output_docx)
    print(output_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
