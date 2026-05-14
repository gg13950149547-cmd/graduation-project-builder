from __future__ import annotations

import argparse
import io
import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree as ET
from docx.shared import RGBColor


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {"w": W_NS, "r": R_NS}


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def iter_runs(paragraph):
    for run in paragraph.runs:
        yield run


def set_run_fonts(run, east_asia: str, ascii_font: str = "Times New Roman", size_pt: float | None = None, bold=None):
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    run.font.name = ascii_font
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), ascii_font)
    color = rpr.color
    if color is None:
        color = OxmlElement("w:color")
        rpr.append(color)
    color.set(qn("w:val"), "000000")
    for attr in (qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade")):
        if attr in color.attrib:
            del color.attrib[attr]


def set_paragraph_metrics(
    paragraph,
    *,
    alignment,
    first_line_indent_pt: float | None,
    space_before_pt: float,
    space_after_pt: float,
    line_spacing,
):
    paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.first_line_indent = None if first_line_indent_pt is None else Pt(first_line_indent_pt)
    pf.left_indent = None
    pf.right_indent = None
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing = line_spacing


def format_all_runs(paragraph, east_asia: str, size_pt: float | None = None, bold=None):
    for run in iter_runs(paragraph):
        if run.text == "":
            continue
        set_run_fonts(run, east_asia=east_asia, size_pt=size_pt, bold=bold)


def append_citation(paragraph, citation: str):
    if citation in paragraph.text:
        return
    paragraph.add_run(citation)
    set_run_fonts(paragraph.runs[-1], east_asia="SimSun", size_pt=12.0, bold=False)


def style_front_matter(doc: Document):
    title_map = {
        "摘要": {"size": 18, "east": "SimHei", "bold": True},
        "Abstract": {"size": 15, "east": "Times New Roman", "bold": True},
        "目录": {"size": 18, "east": "SimHei", "bold": True},
        "本科毕业设计（论文）诚信承诺书": {"size": 18, "east": "SimHei", "bold": True},
    }
    for idx, p in enumerate(doc.paragraphs[:91], start=1):
        if p.text.startswith("Key Words:"):
            p.text = p.text.replace("Key Words:", "KeyWords:", 1)
        text_norm = normalize(p.text)
        if text_norm in title_map:
            cfg = title_map[text_norm]
            set_paragraph_metrics(
                p,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_pt=0,
                space_before_pt=20,
                space_after_pt=20,
                line_spacing=1.0,
            )
            format_all_runs(p, east_asia=cfg["east"], size_pt=cfg["size"], bold=cfg["bold"])
            continue

        if p.text.startswith("关键词："):
            set_paragraph_metrics(
                p,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent_pt=0,
                space_before_pt=3.5,
                space_after_pt=0,
                line_spacing=1.0,
            )
            pf = p.paragraph_format
            pf.left_indent = Pt(26)
            for run in p.runs:
                set_run_fonts(run, east_asia="SimSun" if re.search(r"[\u4e00-\u9fff]", run.text) else "Times New Roman", size_pt=12, bold=True)
            continue

        if p.text.startswith("Key Words:") or p.text.startswith("KeyWords:"):
            set_paragraph_metrics(
                p,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent_pt=0,
                space_before_pt=3.5,
                space_after_pt=0,
                line_spacing=1.0,
            )
            pf = p.paragraph_format
            pf.left_indent = Pt(26)
            for run in p.runs:
                set_run_fonts(run, east_asia="Times New Roman", size_pt=12, bold=True)
            continue

        if 17 <= idx <= 27 and p.text.strip():
            east = "Times New Roman" if re.search(r"[A-Za-z]", p.text) and not re.search(r"[\u4e00-\u9fff]", p.text) else "SimSun"
            is_english_abstract_title = p.text.strip() == "Abstract"
            is_english_abstract_body = 24 <= idx <= 26
            alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            first_indent = 24
            line_spacing = 1.3
            size = 12
            bold = False
            if is_english_abstract_title:
                alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_indent = 0
                line_spacing = 1.0
                size = 15
                bold = True
            elif is_english_abstract_body:
                line_spacing = 1.42
            set_paragraph_metrics(
                p,
                alignment=alignment,
                first_line_indent_pt=first_indent,
                space_before_pt=15.8 if is_english_abstract_title else 3.5 if is_english_abstract_body else 0,
                space_after_pt=0 if is_english_abstract_body else 10 if is_english_abstract_title else 10,
                line_spacing=line_spacing,
            )
            if is_english_abstract_body:
                p.paragraph_format.left_indent = Pt(1.4)
            format_all_runs(p, east_asia=east, size_pt=size, bold=bold)


def style_body(doc: Document):
    ref_titles = {"结论", "结论", "参考文献", "致谢"}
    citation_map = {
        "近年来，随着教育信息化的深入发展和人工智能技术的不断进步": "[10][11]",
        "YOLO系列算法以其高效的实时检测能力在目标检测领域占据重要地位": "[3][5][12]",
        "国内外研究者已经在课堂评估系统方面开展了大量工作": "[1][4][8]",
        "张乐乐等人提出了多模态数据支持的课堂教学行为分析模型与实践框架": "[7][9]",
        "YOLOv8是Ultralytics团队在YOLO系列基础上提出的最新版本": "[12]",
    }

    in_references = False
    for idx, p in enumerate(doc.paragraphs[91:], start=92):
        if p.text == "5.4.2 评估报告生成的设计与与实现":
            p.text = "5.4.2 评估报告生成的设计与实现"
        raw = p.text.strip()
        text_norm = normalize(raw)

        if not raw:
            continue

        if text_norm == "参考文献":
            in_references = True

        if re.match(r"^\d+\s", raw):
            set_paragraph_metrics(
                p,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent_pt=0,
                space_before_pt=8,
                space_after_pt=6,
                line_spacing=1.15,
            )
            format_all_runs(p, east_asia="SimHei", size_pt=16, bold=True)
            continue

        if re.match(r"^\d+\.\d+\s", raw):
            set_paragraph_metrics(
                p,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent_pt=0,
                space_before_pt=6,
                space_after_pt=4,
                line_spacing=1.15,
            )
            format_all_runs(p, east_asia="SimHei", size_pt=14, bold=True)
            continue

        if re.match(r"^\d+\.\d+\.\d+\s", raw):
            set_paragraph_metrics(
                p,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent_pt=0,
                space_before_pt=4,
                space_after_pt=2,
                line_spacing=1.15,
            )
            format_all_runs(p, east_asia="SimHei", size_pt=13, bold=True)
            continue

        if text_norm in {"结论", "致谢", "参考文献"}:
            set_paragraph_metrics(
                p,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_pt=0,
                space_before_pt=10,
                space_after_pt=12,
                line_spacing=1.0,
            )
            format_all_runs(p, east_asia="SimHei", size_pt=18, bold=True)
            continue

        if in_references and raw.startswith("["):
            set_paragraph_metrics(
                p,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent_pt=0,
                space_before_pt=0,
                space_after_pt=2,
                line_spacing=1.3,
            )
            format_all_runs(
                p,
                east_asia="SimSun" if re.search(r"[\u4e00-\u9fff]", raw) else "Times New Roman",
                size_pt=12,
                bold=False,
            )
            continue

        set_paragraph_metrics(
            p,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent_pt=24,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.3,
        )
        for run in iter_runs(p):
            east = "SimSun" if re.search(r"[\u4e00-\u9fff]", run.text) else "Times New Roman"
            set_run_fonts(run, east_asia=east, size_pt=12, bold=False)

        for prefix, citation in citation_map.items():
            if raw.startswith(prefix):
                append_citation(p, citation)
                break


def force_black_all_runs(doc: Document):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            east = "SimSun" if re.search(r"[\u4e00-\u9fff]", run.text) else "Times New Roman"
            if paragraph.text.strip() == "Abstract":
                east = "Times New Roman"
            set_run_fonts(run, east_asia=east, size_pt=run.font.size.pt if run.font.size else None, bold=run.bold)


def set_section_geometry(sect_pr, *, left: str, right: str, top: str, bottom: str, header: str, footer: str):
    pg_sz = sect_pr.find(qn("w:pgSz"))
    if pg_sz is None:
        pg_sz = ET.SubElement(sect_pr, qn("w:pgSz"))
    pg_sz.attrib[qn("w:w")] = "11906"
    pg_sz.attrib[qn("w:h")] = "16838"

    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is None:
        pg_mar = ET.SubElement(sect_pr, qn("w:pgMar"))
    pg_mar.attrib[qn("w:top")] = top
    pg_mar.attrib[qn("w:right")] = right
    pg_mar.attrib[qn("w:bottom")] = bottom
    pg_mar.attrib[qn("w:left")] = left
    pg_mar.attrib[qn("w:header")] = header
    pg_mar.attrib[qn("w:footer")] = footer
    pg_mar.attrib[qn("w:gutter")] = "0"


def ensure_header_footer_refs(sect_pr, header_rid: str, footer_rid: str):
    for tag in (qn("w:headerReference"), qn("w:footerReference")):
        for elem in sect_pr.findall(tag):
            sect_pr.remove(elem)
    hdr = ET.Element(qn("w:headerReference"))
    hdr.attrib[qn("r:id")] = header_rid
    hdr.attrib[qn("w:type")] = "default"
    ftr = ET.Element(qn("w:footerReference"))
    ftr.attrib[qn("r:id")] = footer_rid
    ftr.attrib[qn("w:type")] = "default"
    sect_pr.insert(0, ftr)
    sect_pr.insert(0, hdr)


def set_pg_num_type(sect_pr, *, fmt: str | None, start: str | None):
    elem = sect_pr.find(qn("w:pgNumType"))
    if elem is None:
        elem = ET.SubElement(sect_pr, qn("w:pgNumType"))
    if fmt is not None:
        elem.attrib[qn("w:fmt")] = fmt
    elif qn("w:fmt") in elem.attrib:
        del elem.attrib[qn("w:fmt")]
    if start is not None:
        elem.attrib[qn("w:start")] = start
    elif qn("w:start") in elem.attrib:
        del elem.attrib[qn("w:start")]


def patch_section_xml(docx_path: Path):
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_dir_path = Path(tmp_dir)
        with zipfile.ZipFile(docx_path) as zin:
            zin.extractall(tmp_dir_path)

        doc_xml = tmp_dir_path / "word" / "document.xml"
        rels_xml = tmp_dir_path / "word" / "_rels" / "document.xml.rels"
        header_xml = tmp_dir_path / "word" / "header1.xml"
        footer_xml = tmp_dir_path / "word" / "footer1.xml"

        rels_root = ET.parse(rels_xml).getroot()
        relmap = {rel.get("Id"): rel.get("Target") for rel in rels_root}

        tree = ET.parse(doc_xml)
        root = tree.getroot()
        sects = root.xpath("//w:sectPr", namespaces=NS)
        if len(sects) != 9:
            raise RuntimeError(f"Unexpected section count: {len(sects)}")

        body_header_rid = None
        body_footer_rid = None
        for ref in sects[5].xpath("./w:headerReference/@r:id", namespaces=NS):
            body_header_rid = ref
        for ref in sects[5].xpath("./w:footerReference/@r:id", namespaces=NS):
            body_footer_rid = ref
        if not body_header_rid or not body_footer_rid:
            raise RuntimeError("Could not locate body header/footer relationship ids")

        # Front matter: Chinese abstract, English abstract, TOC.
        for idx in (2, 3, 4):
            ensure_header_footer_refs(sects[idx], body_header_rid, body_footer_rid)
            set_section_geometry(
                sects[idx],
                left="1673",
                right="1235",
                top="1501",
                bottom="1332",
                header="1105",
                footer="1129",
            )
            if idx == 2:
                set_pg_num_type(sects[idx], fmt="upperRoman", start="1")
            else:
                set_pg_num_type(sects[idx], fmt="upperRoman", start=None)

        # Body and tail blocks: body, conclusion, acknowledgement, references, final body sectPr.
        for idx in (5, 6, 7, 8):
            ensure_header_footer_refs(sects[idx], body_header_rid, body_footer_rid)
            set_section_geometry(
                sects[idx],
                left="1673",
                right="1235",
                top="1501",
                bottom="1332",
                header="1105",
                footer="1129",
            )
        set_pg_num_type(sects[5], fmt="decimal", start="1")
        for idx in (6, 7, 8):
            set_pg_num_type(sects[idx], fmt="decimal", start=None)

        tree.write(doc_xml, encoding="utf-8", xml_declaration=True)

        # Darken header/footer to match the reference template more closely.
        for xml_path in (header_xml, footer_xml):
            xml_tree = ET.parse(xml_path)
            xml_root = xml_tree.getroot()
            for color in xml_root.xpath(".//w:color", namespaces=NS):
                color.attrib[qn("w:val")] = "000000"
            for border in xml_root.xpath(".//w:bottom", namespaces=NS):
                border.attrib[qn("w:color")] = "000000"
                border.attrib[qn("w:sz")] = "8"
            xml_tree.write(xml_path, encoding="utf-8", xml_declaration=True)

        tmp_docx = docx_path.with_suffix(".patched.docx")
        with zipfile.ZipFile(tmp_docx, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for file_path in sorted(tmp_dir_path.rglob("*")):
                if file_path.is_file():
                    zout.write(file_path, file_path.relative_to(tmp_dir_path).as_posix())
        shutil.move(tmp_docx, docx_path)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(input_path, output_path)

    doc = Document(output_path)
    style_front_matter(doc)
    style_body(doc)
    force_black_all_runs(doc)
    doc.save(output_path)

    patch_section_xml(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
